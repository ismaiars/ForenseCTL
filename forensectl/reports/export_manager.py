"""Gestor de exportación de reportes a diferentes formatos."""

import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional

from forensectl import logger


class ExportManager:
    """Gestor de exportación de reportes a múltiples formatos."""
    
    def __init__(self) -> None:
        """Inicializar gestor de exportación."""
        # Configuración de herramientas de conversión
        self.conversion_tools = {
            "pdf": {
                "tool": "wkhtmltopdf",
                "fallback": "weasyprint",
                "options": [
                    "--page-size", "A4",
                    "--margin-top", "0.75in",
                    "--margin-right", "0.75in",
                    "--margin-bottom", "0.75in",
                    "--margin-left", "0.75in",
                    "--encoding", "UTF-8",
                    "--enable-local-file-access"
                ]
            },
            "docx": {
                "tool": "pandoc",
                "options": [
                    "-f", "html",
                    "-t", "docx"
                ]
            }
        }
        
        # Verificar disponibilidad de herramientas
        self.available_tools = self._check_available_tools()
    
    def export_report(
        self,
        content: str,
        output_dir: Path,
        output_format: str,
        metadata: Dict[str, Any]
    ) -> Path:
        """Exportar reporte al formato especificado.
        
        Args:
            content: Contenido renderizado del reporte
            output_dir: Directorio de salida
            output_format: Formato de salida (pdf, html, markdown, docx)
            metadata: Metadatos del reporte
            
        Returns:
            Ruta del archivo exportado
        """
        logger.info(f"Exportando reporte a formato {output_format}")
        
        # Generar nombre de archivo
        case_id = metadata.get("case_id", "unknown")
        timestamp = metadata.get("generated_at", "unknown").replace(":", "-")
        filename = f"report_{case_id}_{timestamp}.{output_format}"
        output_file = output_dir / filename
        
        try:
            if output_format == "html":
                return self._export_html(content, output_file)
            elif output_format == "pdf":
                return self._export_pdf(content, output_file, metadata)
            elif output_format == "markdown":
                return self._export_markdown(content, output_file)
            elif output_format == "docx":
                return self._export_docx(content, output_file, metadata)
            else:
                raise ValueError(f"Formato no soportado: {output_format}")
                
        except Exception as e:
            logger.error(f"Error exportando a {output_format}: {e}")
            # Fallback a HTML
            if output_format != "html":
                logger.info("Usando fallback a HTML")
                fallback_file = output_dir / f"report_{case_id}_{timestamp}.html"
                return self._export_html(content, fallback_file)
            raise
    
    def _export_html(self, content: str, output_file: Path) -> Path:
        """Exportar a HTML.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            
        Returns:
            Ruta del archivo exportado
        """
        try:
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Reporte HTML exportado: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Error exportando HTML: {e}")
            raise
    
    def _export_pdf(self, content: str, output_file: Path, metadata: Dict[str, Any]) -> Path:
        """Exportar a PDF.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            metadata: Metadatos del reporte
            
        Returns:
            Ruta del archivo exportado
        """
        try:
            # Intentar con wkhtmltopdf primero
            if "wkhtmltopdf" in self.available_tools:
                return self._export_pdf_wkhtmltopdf(content, output_file, metadata)
            
            # Fallback a weasyprint
            elif "weasyprint" in self.available_tools:
                return self._export_pdf_weasyprint(content, output_file)
            
            # Fallback a conversión básica (simulada)
            else:
                logger.warning("No hay herramientas PDF disponibles, usando fallback")
                return self._export_pdf_fallback(content, output_file)
                
        except Exception as e:
            logger.error(f"Error exportando PDF: {e}")
            raise
    
    def _export_pdf_wkhtmltopdf(
        self,
        content: str,
        output_file: Path,
        metadata: Dict[str, Any]
    ) -> Path:
        """Exportar PDF usando wkhtmltopdf.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            metadata: Metadatos del reporte
            
        Returns:
            Ruta del archivo exportado
        """
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as temp_html:
            temp_html.write(content)
            temp_html_path = temp_html.name
        
        try:
            # Construir comando
            cmd = ["wkhtmltopdf"]
            cmd.extend(self.conversion_tools["pdf"]["options"])
            
            # Agregar metadatos si están disponibles
            if metadata.get("title"):
                cmd.extend(["--title", metadata["title"]])
            
            cmd.extend([temp_html_path, str(output_file)])
            
            # Ejecutar conversión
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300  # 5 minutos timeout
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"wkhtmltopdf falló: {result.stderr}")
            
            logger.info(f"PDF generado con wkhtmltopdf: {output_file}")
            return output_file
            
        finally:
            # Limpiar archivo temporal
            Path(temp_html_path).unlink(missing_ok=True)
    
    def _export_pdf_weasyprint(self, content: str, output_file: Path) -> Path:
        """Exportar PDF usando WeasyPrint.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            
        Returns:
            Ruta del archivo exportado
        """
        try:
            import weasyprint
            
            # Generar PDF
            html_doc = weasyprint.HTML(string=content)
            html_doc.write_pdf(str(output_file))
            
            logger.info(f"PDF generado con WeasyPrint: {output_file}")
            return output_file
            
        except ImportError:
            logger.error("WeasyPrint no está instalado")
            raise
        except Exception as e:
            logger.error(f"Error con WeasyPrint: {e}")
            raise
    
    def _export_pdf_fallback(self, content: str, output_file: Path) -> Path:
        """Exportar PDF usando método de fallback.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            
        Returns:
            Ruta del archivo exportado
        """
        # Como fallback, guardar como HTML con extensión PDF
        # En un entorno real, se podría usar una biblioteca como reportlab
        logger.warning("Usando fallback básico para PDF (guardando como HTML)")
        
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(content)
        
        return output_file
    
    def _export_markdown(self, content: str, output_file: Path) -> Path:
        """Exportar a Markdown.
        
        Args:
            content: Contenido (puede ser HTML o Markdown)
            output_file: Archivo de salida
            
        Returns:
            Ruta del archivo exportado
        """
        try:
            # Si el contenido es HTML, convertir a Markdown
            if content.strip().startswith("<"):
                # Intentar conversión con pandoc
                if "pandoc" in self.available_tools:
                    content = self._convert_html_to_markdown_pandoc(content)
                else:
                    # Conversión básica
                    content = self._convert_html_to_markdown_basic(content)
            
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(content)
            
            logger.info(f"Reporte Markdown exportado: {output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Error exportando Markdown: {e}")
            raise
    
    def _export_docx(self, content: str, output_file: Path, metadata: Dict[str, Any]) -> Path:
        """Exportar a DOCX.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            metadata: Metadatos del reporte
            
        Returns:
            Ruta del archivo exportado
        """
        try:
            if "pandoc" in self.available_tools:
                return self._export_docx_pandoc(content, output_file)
            else:
                # Fallback usando python-docx
                return self._export_docx_python_docx(content, output_file, metadata)
                
        except Exception as e:
            logger.error(f"Error exportando DOCX: {e}")
            raise
    
    def _export_docx_pandoc(self, content: str, output_file: Path) -> Path:
        """Exportar DOCX usando Pandoc.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            
        Returns:
            Ruta del archivo exportado
        """
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as temp_html:
            temp_html.write(content)
            temp_html_path = temp_html.name
        
        try:
            # Construir comando pandoc
            cmd = ["pandoc"]
            cmd.extend(self.conversion_tools["docx"]["options"])
            cmd.extend(["-o", str(output_file), temp_html_path])
            
            # Ejecutar conversión
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=300
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"Pandoc falló: {result.stderr}")
            
            logger.info(f"DOCX generado con Pandoc: {output_file}")
            return output_file
            
        finally:
            Path(temp_html_path).unlink(missing_ok=True)
    
    def _export_docx_python_docx(self, content: str, output_file: Path, metadata: Dict[str, Any]) -> Path:
        """Exportar DOCX usando python-docx.
        
        Args:
            content: Contenido HTML
            output_file: Archivo de salida
            metadata: Metadatos del reporte
            
        Returns:
            Ruta del archivo exportado
        """
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            
            # Crear documento
            doc = Document()
            
            # Agregar título
            title = metadata.get("title", "Reporte Forense")
            doc.add_heading(title, 0)
            
            # Parsear HTML y convertir a contenido DOCX
            soup = BeautifulSoup(content, 'html.parser')
            
            # Procesar elementos HTML
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table']):
                if element.name in ['h1', 'h2', 'h3']:
                    level = int(element.name[1])
                    doc.add_heading(element.get_text(), level)
                elif element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name == 'table':
                    # Agregar tabla básica
                    rows = element.find_all('tr')
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                        for i, row in enumerate(rows):
                            cells = row.find_all(['td', 'th'])
                            for j, cell in enumerate(cells):
                                if i < len(table.rows) and j < len(table.rows[i].cells):
                                    table.rows[i].cells[j].text = cell.get_text()
            
            # Guardar documento
            doc.save(str(output_file))
            
            logger.info(f"DOCX generado con python-docx: {output_file}")
            return output_file
            
        except ImportError:
            logger.error("python-docx o beautifulsoup4 no están instalados")
            raise
        except Exception as e:
            logger.error(f"Error con python-docx: {e}")
            raise
    
    def _convert_html_to_markdown_pandoc(self, html_content: str) -> str:
        """Convertir HTML a Markdown usando Pandoc.
        
        Args:
            html_content: Contenido HTML
            
        Returns:
            Contenido en Markdown
        """
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False, encoding="utf-8") as temp_html:
            temp_html.write(html_content)
            temp_html_path = temp_html.name
        
        try:
            cmd = ["pandoc", "-f", "html", "-t", "markdown", temp_html_path]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"Pandoc falló: {result.stderr}")
            
            return result.stdout
            
        finally:
            Path(temp_html_path).unlink(missing_ok=True)
    
    def _convert_html_to_markdown_basic(self, html_content: str) -> str:
        """Conversión básica de HTML a Markdown.
        
        Args:
            html_content: Contenido HTML
            
        Returns:
            Contenido en Markdown
        """
        try:
            from bs4 import BeautifulSoup
            import re
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Conversiones básicas
            markdown_content = []
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    markdown_content.append('#' * level + ' ' + element.get_text().strip())
                elif element.name == 'p':
                    text = element.get_text().strip()
                    if text:
                        markdown_content.append(text)
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li'):
                        prefix = '- ' if element.name == 'ul' else '1. '
                        markdown_content.append(prefix + li.get_text().strip())
                elif element.name == 'table':
                    # Conversión básica de tabla
                    rows = element.find_all('tr')
                    if rows:
                        # Encabezados
                        headers = rows[0].find_all(['th', 'td'])
                        if headers:
                            header_row = '| ' + ' | '.join(h.get_text().strip() for h in headers) + ' |'
                            separator = '|' + '---|' * len(headers)
                            markdown_content.extend([header_row, separator])
                            
                            # Filas de datos
                            for row in rows[1:]:
                                cells = row.find_all(['td', 'th'])
                                if cells:
                                    data_row = '| ' + ' | '.join(c.get_text().strip() for c in cells) + ' |'
                                    markdown_content.append(data_row)
                
                markdown_content.append('')  # Línea en blanco
            
            return '\n'.join(markdown_content)
            
        except ImportError:
            logger.warning("BeautifulSoup no disponible, usando conversión muy básica")
            # Conversión muy básica usando regex
            content = html_content
            content = re.sub(r'<h([1-6]).*?>(.*?)</h[1-6]>', lambda m: '#' * int(m.group(1)) + ' ' + m.group(2), content)
            content = re.sub(r'<p.*?>(.*?)</p>', r'\1\n', content)
            content = re.sub(r'<[^>]+>', '', content)  # Remover todas las etiquetas HTML
            return content
    
    def _check_available_tools(self) -> Dict[str, bool]:
        """Verificar qué herramientas de conversión están disponibles.
        
        Returns:
            Diccionario con disponibilidad de herramientas
        """
        tools = {}
        
        # Verificar wkhtmltopdf
        try:
            result = subprocess.run(
                ["wkhtmltopdf", "--version"],
                capture_output=True,
                timeout=10
            )
            tools["wkhtmltopdf"] = result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError):
            tools["wkhtmltopdf"] = False
        
        # Verificar pandoc
        try:
            result = subprocess.run(
                ["pandoc", "--version"],
                capture_output=True,
                timeout=10
            )
            tools["pandoc"] = result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError):
            tools["pandoc"] = False
        
        # Verificar WeasyPrint
        try:
            import weasyprint
            tools["weasyprint"] = True
        except ImportError:
            tools["weasyprint"] = False
        
        # Verificar python-docx
        try:
            import docx
            tools["python-docx"] = True
        except ImportError:
            tools["python-docx"] = False
        
        # Verificar BeautifulSoup
        try:
            import bs4
            tools["beautifulsoup4"] = True
        except ImportError:
            tools["beautifulsoup4"] = False
        
        logger.info(f"Herramientas de conversión disponibles: {tools}")
        return tools
    
    def get_supported_formats(self) -> List[str]:
        """Obtener lista de formatos soportados.
        
        Returns:
            Lista de formatos soportados
        """
        supported = ["html", "markdown"]
        
        if self.available_tools.get("wkhtmltopdf") or self.available_tools.get("weasyprint"):
            supported.append("pdf")
        
        if self.available_tools.get("pandoc") or self.available_tools.get("python-docx"):
            supported.append("docx")
        
        return supported
    
    def get_conversion_info(self) -> Dict[str, Any]:
        """Obtener información sobre capacidades de conversión.
        
        Returns:
            Información sobre herramientas y capacidades
        """
        return {
            "available_tools": self.available_tools,
            "supported_formats": self.get_supported_formats(),
            "recommended_tools": {
                "pdf": "wkhtmltopdf" if self.available_tools.get("wkhtmltopdf") else "weasyprint",
                "docx": "pandoc" if self.available_tools.get("pandoc") else "python-docx",
                "markdown": "pandoc" if self.available_tools.get("pandoc") else "basic"
            }
        }