#!/usr/bin/env python3
"""
Demo Completo Interactivo de ForenseCTL

Este script demuestra TODAS las funcionalidades principales de ForenseCTL:
- Gestión completa de casos
- Recopilación de evidencia
- Análisis forense
- Generación de reportes en múltiples formatos
- Cadena de custodia
- Exportación de datos
- Interfaz interactiva para el usuario
"""

import os
import sys
import json
import subprocess
from datetime import datetime
from pathlib import Path

# Agregar el directorio del proyecto al path
sys.path.insert(0, str(Path(__file__).parent))

from forensectl.core.case_manager import CaseManager
from forensectl.core.chain_of_custody import ChainOfCustody
from forensectl.core.manifest import Manifest
from forensectl.analysis.artifact_extractor import ArtifactExtractor
from forensectl.analysis.timeline_builder import TimelineBuilder
from forensectl.reports.report_generator import ReportGenerator
from forensectl.reports.export_manager import ExportManager
from real_system_analyzer import RealSystemAnalyzer


class ForenseCTLDemo:
    """Clase principal para el demo interactivo de ForenseCTL."""
    
    def __init__(self):
        self.case_manager = None
        self.current_case_id = None
        self.chain_of_custody = None
        self.manifest = None
        
    def mostrar_banner(self):
        """Mostrar banner principal."""
        print("\n" + "="*80)
        print("🚀 FORENSECTL - DEMO COMPLETO INTERACTIVO")
        print("="*80)
        print("🔍 Framework Completo de Análisis Forense Digital")
        print("📋 Todas las funcionalidades disponibles")
        print("🎯 Interfaz interactiva para manejo completo")
        print("="*80 + "\n")
        
    def mostrar_menu_principal(self):
        """Mostrar menú principal de opciones."""
        print("\n📋 MENÚ PRINCIPAL - FORENSECTL")
        print("-" * 50)
        print("1. 📁 Gestión de Casos")
        print("2. 🔍 Recopilación de Evidencia")
        print("3. ⚙️ Análisis Forense")
        print("4. 📄 Generación de Reportes")
        print("5. 🔗 Cadena de Custodia")
        print("6. 📊 Exportación de Datos")
        print("7. 📈 Estado del Sistema")
        print("8. 🛠️ Herramientas Adicionales")
        print("9. 🔬 Análisis Forense Completo del Sistema")
        print("10. ❓ Ayuda y Documentación")
        print("0. 🚪 Salir")
        print("-" * 50)
        
    def gestionar_casos(self):
        """Submenu para gestión de casos."""
        while True:
            print("\n📁 GESTIÓN DE CASOS")
            print("-" * 30)
            print("1. ➕ Crear nuevo caso")
            print("2. 📋 Listar casos existentes")
            print("3. 🔍 Seleccionar caso activo")
            print("4. ℹ️ Ver información del caso")
            print("5. 📝 Actualizar información")
            print("6. 🗑️ Eliminar caso")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.crear_caso()
            elif opcion == "2":
                self.listar_casos()
            elif opcion == "3":
                self.seleccionar_caso()
            elif opcion == "4":
                self.ver_info_caso()
            elif opcion == "5":
                self.actualizar_caso()
            elif opcion == "6":
                self.eliminar_caso()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def crear_caso(self):
        """Crear un nuevo caso."""
        print("\n➕ CREAR NUEVO CASO")
        print("-" * 25)
        
        # Solicitar información del caso
        examiner = input("👤 Nombre del examinador: ").strip() or "Analista Forense"
        organization = input("🏢 Organización: ").strip() or "Centro de Ciberseguridad"
        description = input("📝 Descripción del caso: ").strip() or "Análisis forense digital"
        
        # Generar ID único y simple
        if not self.case_manager:
            self.case_manager = CaseManager()
        
        # Obtener número de casos existentes para generar ID secuencial
        existing_cases = self.case_manager.list_cases()
        case_number = len(existing_cases) + 1
        case_id = f"CASO-{case_number:03d}"
        
        try:
            
            # Crear caso
            case_info = self.case_manager.create_case(
                case_id=case_id,
                examiner=examiner,
                organization=organization,
                description=description,
                timezone_str="UTC"
            )
            
            # Configurar estructura
            self.case_manager.setup_case_structure(case_id)
            
            # Establecer como caso activo
            self.current_case_id = case_id
            self.chain_of_custody = ChainOfCustody(case_id)
            self.manifest = Manifest(case_id)
            
            print(f"\n✅ Caso creado exitosamente: {case_id}")
            print(f"📂 Directorio: cases/{case_id}")
            print(f"👤 Examinador: {examiner}")
            print(f"🏢 Organización: {organization}")
            
        except Exception as e:
            print(f"❌ Error creando caso: {e}")
            
    def listar_casos(self):
        """Listar todos los casos existentes."""
        print("\n📋 CASOS EXISTENTES")
        print("-" * 25)
        
        try:
            if not self.case_manager:
                self.case_manager = CaseManager()
                
            casos = self.case_manager.list_cases()
            
            if not casos:
                print("📭 No hay casos creados")
                return
                
            for i, caso in enumerate(casos, 1):
                status_icon = "🟢" if caso.get('status') == 'active' else "🔴"
                print(f"{i}. {status_icon} {caso['case_id']}")
                print(f"   👤 {caso.get('examiner', 'N/A')}")
                print(f"   📅 {caso.get('created_at', 'N/A')[:10]}")
                print(f"   📝 {caso.get('description', 'N/A')[:50]}...")
                print()
                
        except Exception as e:
            print(f"❌ Error listando casos: {e}")
            
    def seleccionar_caso(self):
        """Seleccionar un caso como activo."""
        print("\n🔍 SELECCIONAR CASO ACTIVO")
        print("-" * 30)
        
        case_id = input("📁 ID del caso: ").strip()
        
        if not case_id:
            print("❌ ID de caso requerido")
            return
            
        try:
            if not self.case_manager:
                self.case_manager = CaseManager()
                
            # Verificar que el caso existe
            case_info = self.case_manager.get_case_info(case_id)
            
            if case_info is None:
                print(f"❌ No se pudo encontrar el caso {case_id}")
                return
            
            # Establecer como caso activo
            self.current_case_id = case_id
            self.chain_of_custody = ChainOfCustody(case_id)
            self.manifest = Manifest(case_id)
            
            print(f"✅ Caso activo: {case_id}")
            print(f"👤 Examinador: {case_info.get('examiner', 'N/A')}")
            print(f"📅 Creado: {case_info.get('created_at', 'N/A')[:10]}")
            
        except Exception as e:
            print(f"❌ Error seleccionando caso: {e}")
            
    def ver_info_caso(self):
        """Ver información detallada del caso activo."""
        if not self.current_case_id:
            print("❌ No hay caso activo seleccionado")
            return
            
        print(f"\nℹ️ INFORMACIÓN DEL CASO: {self.current_case_id}")
        print("-" * 50)
        
        try:
            case_info = self.case_manager.get_case_info(self.current_case_id)
            
            if case_info is None:
                print(f"❌ No se pudo encontrar información del caso {self.current_case_id}")
                return
            
            print(f"📁 ID: {case_info.get('case_id', 'N/A')}")
            print(f"🆔 UUID: {case_info.get('uuid', 'N/A')}")
            print(f"👤 Examinador: {case_info.get('examiner', 'N/A')}")
            print(f"🏢 Organización: {case_info.get('organization', 'N/A')}")
            print(f"📝 Descripción: {case_info.get('description', 'N/A')}")
            print(f"🌍 Zona horaria: {case_info.get('timezone', 'N/A')}")
            print(f"📅 Creado: {case_info.get('created_at', 'N/A')}")
            print(f"🔄 Estado: {case_info.get('status', 'N/A')}")
            print(f"📊 Versión: {case_info.get('version', 'N/A')}")
            
        except Exception as e:
            print(f"❌ Error obteniendo información: {e}")
            
    def actualizar_caso(self):
        """Actualizar información del caso."""
        if not self.current_case_id:
            print("❌ No hay caso activo seleccionado")
            return
            
        print(f"\n📝 ACTUALIZAR CASO: {self.current_case_id}")
        print("-" * 40)
        
        try:
            case_info = self.case_manager.get_case_info(self.current_case_id)
            
            if case_info is None:
                print(f"❌ No se pudo encontrar información del caso {self.current_case_id}")
                return
            
            print("Presiona Enter para mantener el valor actual:")
            
            new_examiner = input(f"👤 Examinador [{case_info.get('examiner', '')}]: ").strip()
            new_org = input(f"🏢 Organización [{case_info.get('organization', '')}]: ").strip()
            new_desc = input(f"📝 Descripción [{case_info.get('description', '')}]: ").strip()
            
            updates = {}
            if new_examiner:
                updates['examiner'] = new_examiner
            if new_org:
                updates['organization'] = new_org
            if new_desc:
                updates['description'] = new_desc
                
            if updates:
                self.case_manager.update_case_info(self.current_case_id, updates)
                print("✅ Caso actualizado exitosamente")
            else:
                print("ℹ️ No se realizaron cambios")
                
        except Exception as e:
            print(f"❌ Error actualizando caso: {e}")
            
    def eliminar_caso(self):
        """Eliminar un caso."""
        print("\n🗑️ ELIMINAR CASO")
        print("-" * 20)
        
        case_id = input("📁 ID del caso a eliminar: ").strip()
        
        if not case_id:
            print("❌ ID de caso requerido")
            return
            
        confirmacion = input(f"⚠️ ¿Estás seguro de eliminar {case_id}? (sí/no): ").strip().lower()
        
        if confirmacion in ['sí', 'si', 'yes', 'y']:
            try:
                success = self.case_manager.delete_case(case_id, confirm=True)
                
                if success:
                    if self.current_case_id == case_id:
                        self.current_case_id = None
                        self.chain_of_custody = None
                        self.manifest = None
                        
                    print(f"✅ Caso {case_id} eliminado")
                else:
                    print(f"❌ No se pudo eliminar el caso {case_id}")
                
            except Exception as e:
                print(f"❌ Error eliminando caso: {e}")
        else:
            print("❌ Eliminación cancelada")
            
    def recopilar_evidencia(self):
        """Submenu para recopilación de evidencia."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        while True:
            print(f"\n🔍 RECOPILACIÓN DE EVIDENCIA - {self.current_case_id}")
            print("-" * 50)
            print("1. 💻 Información del sistema")
            print("2. ⚙️ Procesos activos")
            print("3. 🌐 Conexiones de red")
            print("4. 📁 Archivos temporales")
            print("5. 🌍 Variables de entorno")
            print("6. 📊 Recopilación automática completa")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.recopilar_info_sistema()
            elif opcion == "2":
                self.recopilar_procesos()
            elif opcion == "3":
                self.recopilar_red()
            elif opcion == "4":
                self.recopilar_archivos_temp()
            elif opcion == "5":
                self.recopilar_variables_entorno()
            elif opcion == "6":
                self.recopilacion_automatica()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def recopilar_info_sistema(self):
        """Recopilar información básica del sistema."""
        print("\n💻 RECOPILANDO INFORMACIÓN DEL SISTEMA...")
        
        try:
            system_info = {
                "hostname": os.environ.get("COMPUTERNAME", "unknown"),
                "username": os.environ.get("USERNAME", "unknown"),
                "os_version": sys.platform,
                "python_version": sys.version.split()[0],
                "architecture": os.environ.get("PROCESSOR_ARCHITECTURE", "unknown"),
                "timestamp": datetime.now().isoformat()
            }
            
            # Guardar evidencia
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/system_info.json")
            with open(evidence_file, 'w', encoding='utf-8') as f:
                json.dump(system_info, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            evidence_id = f"SYS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_evidence(
                evidence_id=evidence_id,
                evidence_type="system_info",
                source_path=str(evidence_file),
                description="Información básica del sistema"
            )
            
            print("✅ Información del sistema recopilada")
            print(f"📁 Archivo: {evidence_file}")
            print(f"🆔 ID de evidencia: {evidence_id}")
            
            for key, value in system_info.items():
                print(f"  {key}: {value}")
                
        except Exception as e:
            print(f"❌ Error recopilando información: {e}")
            
    def recopilar_procesos(self):
        """Recopilar información de procesos activos."""
        print("\n⚙️ RECOPILANDO PROCESOS ACTIVOS...")
        
        try:
            # Simular datos de procesos (en un caso real usarías psutil)
            procesos = [
                {"pid": 1234, "name": "explorer.exe", "cpu_percent": 2.5, "memory_percent": 15.2},
                {"pid": 5678, "name": "chrome.exe", "cpu_percent": 8.1, "memory_percent": 25.7},
                {"pid": 9012, "name": "python.exe", "cpu_percent": 12.3, "memory_percent": 8.9},
                {"pid": 3456, "name": "notepad.exe", "cpu_percent": 0.1, "memory_percent": 2.1},
                {"pid": 7890, "name": "winlogon.exe", "cpu_percent": 0.5, "memory_percent": 3.4}
            ]
            
            # Guardar evidencia
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/running_processes.json")
            with open(evidence_file, 'w', encoding='utf-8') as f:
                json.dump(procesos, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            evidence_id = f"PROC-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_evidence(
                evidence_id=evidence_id,
                evidence_type="processes",
                source_path=str(evidence_file),
                description="Procesos activos del sistema"
            )
            
            print("✅ Procesos activos recopilados")
            print(f"📁 Archivo: {evidence_file}")
            print(f"🆔 ID de evidencia: {evidence_id}")
            print(f"📊 Total de procesos: {len(procesos)}")
            
        except Exception as e:
            print(f"❌ Error recopilando procesos: {e}")
            
    def recopilar_red(self):
        """Recopilar información de conexiones de red."""
        print("\n🌐 RECOPILANDO CONEXIONES DE RED...")
        
        try:
            # Simular conexiones de red
            conexiones = [
                {"local_address": "127.0.0.1:8080", "remote_address": "0.0.0.0:0", "status": "LISTENING", "process": "python.exe"},
                {"local_address": "192.168.1.100:443", "remote_address": "8.8.8.8:443", "status": "ESTABLISHED", "process": "chrome.exe"},
                {"local_address": "192.168.1.100:80", "remote_address": "172.217.14.142:80", "status": "ESTABLISHED", "process": "chrome.exe"}
            ]
            
            # Guardar evidencia
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/network_connections.json")
            with open(evidence_file, 'w', encoding='utf-8') as f:
                json.dump(conexiones, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            evidence_id = f"NET-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_evidence(
                evidence_id=evidence_id,
                evidence_type="network",
                source_path=str(evidence_file),
                description="Conexiones de red activas"
            )
            
            print("✅ Conexiones de red recopiladas")
            print(f"📁 Archivo: {evidence_file}")
            print(f"🆔 ID de evidencia: {evidence_id}")
            print(f"🌐 Total de conexiones: {len(conexiones)}")
            
        except Exception as e:
            print(f"❌ Error recopilando conexiones: {e}")
            
    def recopilar_archivos_temp(self):
        """Recopilar información de archivos temporales."""
        print("\n📁 RECOPILANDO ARCHIVOS TEMPORALES...")
        
        try:
            temp_files = []
            temp_dir = Path(os.environ.get("TEMP", "/tmp"))
            
            if temp_dir.exists():
                for file in temp_dir.iterdir():
                    if file.is_file() and len(temp_files) < 20:
                        try:
                            temp_files.append({
                                "name": file.name,
                                "size": file.stat().st_size,
                                "modified": datetime.fromtimestamp(file.stat().st_mtime).isoformat(),
                                "path": str(file)
                            })
                        except (OSError, PermissionError):
                            continue
                            
            # Guardar evidencia
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/temp_files.json")
            with open(evidence_file, 'w', encoding='utf-8') as f:
                json.dump(temp_files, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            evidence_id = f"TEMP-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_evidence(
                evidence_id=evidence_id,
                evidence_type="temp_files",
                source_path=str(evidence_file),
                description="Archivos temporales del sistema"
            )
            
            print("✅ Archivos temporales recopilados")
            print(f"📁 Archivo: {evidence_file}")
            print(f"🆔 ID de evidencia: {evidence_id}")
            print(f"📊 Total de archivos: {len(temp_files)}")
            
        except Exception as e:
            print(f"❌ Error recopilando archivos temporales: {e}")
            
    def recopilar_variables_entorno(self):
        """Recopilar variables de entorno relevantes."""
        print("\n🌍 RECOPILANDO VARIABLES DE ENTORNO...")
        
        try:
            # Variables relevantes para análisis forense
            variables_relevantes = [
                'PATH', 'COMPUTERNAME', 'USERNAME', 'OS', 'PROCESSOR_ARCHITECTURE',
                'TEMP', 'USERPROFILE', 'PROGRAMFILES', 'SYSTEMROOT', 'WINDIR'
            ]
            
            env_vars = {k: os.environ.get(k, 'N/A') for k in variables_relevantes}
            
            # Guardar evidencia
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/environment_vars.json")
            with open(evidence_file, 'w', encoding='utf-8') as f:
                json.dump(env_vars, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            evidence_id = f"ENV-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_evidence(
                evidence_id=evidence_id,
                evidence_type="environment",
                source_path=str(evidence_file),
                description="Variables de entorno del sistema"
            )
            
            print("✅ Variables de entorno recopiladas")
            print(f"📁 Archivo: {evidence_file}")
            print(f"🆔 ID de evidencia: {evidence_id}")
            print(f"🌍 Total de variables: {len(env_vars)}")
            
        except Exception as e:
            print(f"❌ Error recopilando variables: {e}")
            
    def recopilacion_automatica(self):
        """Realizar recopilación automática completa."""
        print("\n📊 RECOPILACIÓN AUTOMÁTICA COMPLETA...")
        print("Esto puede tomar unos momentos...\n")
        
        try:
            # Ejecutar todas las recopilaciones
            self.recopilar_info_sistema()
            print()
            self.recopilar_procesos()
            print()
            self.recopilar_red()
            print()
            self.recopilar_archivos_temp()
            print()
            self.recopilar_variables_entorno()
            
            print("\n🎉 RECOPILACIÓN AUTOMÁTICA COMPLETADA")
            print("✅ Toda la evidencia ha sido recopilada y registrada")
            
        except Exception as e:
            print(f"❌ Error en recopilación automática: {e}")
            
    def analisis_forense(self):
        """Submenu para análisis forense."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        while True:
            print(f"\n⚙️ ANÁLISIS FORENSE - {self.current_case_id}")
            print("-" * 40)
            print("1. 📊 Análisis de procesos")
            print("2. 🌐 Análisis de red")
            print("3. 📁 Análisis de archivos")
            print("4. ⏱️ Construcción de timeline")
            print("5. 🔍 Búsqueda de artefactos")
            print("6. 🛡️ Análisis de seguridad")
            print("7. 📈 Análisis completo automático")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.analizar_procesos()
            elif opcion == "2":
                self.analizar_red()
            elif opcion == "3":
                self.analizar_archivos()
            elif opcion == "4":
                self.construir_timeline()
            elif opcion == "5":
                self.buscar_artefactos()
            elif opcion == "6":
                self.analisis_seguridad()
            elif opcion == "7":
                self.analisis_completo()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def analizar_procesos(self):
        """Analizar procesos del sistema."""
        print("\n📊 ANALIZANDO PROCESOS...")
        
        try:
            # Cargar evidencia de procesos
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/running_processes.json")
            
            if not evidence_file.exists():
                print("❌ No se encontró evidencia de procesos. Ejecuta la recopilación primero.")
                return
                
            with open(evidence_file, 'r', encoding='utf-8') as f:
                procesos = json.load(f)
                
            # Análisis básico
            total_procesos = len(procesos)
            procesos_alto_cpu = [p for p in procesos if p['cpu_percent'] > 10]
            procesos_alta_memoria = [p for p in procesos if p['memory_percent'] > 20]
            
            analisis = {
                "total_procesos": total_procesos,
                "procesos_alto_cpu": procesos_alto_cpu,
                "procesos_alta_memoria": procesos_alta_memoria,
                "timestamp": datetime.now().isoformat(),
                "resumen": {
                    "procesos_sospechosos": len(procesos_alto_cpu),
                    "uso_memoria_alto": len(procesos_alta_memoria),
                    "nivel_riesgo": "BAJO" if len(procesos_alto_cpu) < 2 else "MEDIO"
                }
            }
            
            # Guardar análisis
            analysis_file = Path(f"cases/{self.current_case_id}/analysis/process_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analisis, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            analysis_id = f"PROC-ANALYSIS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_analysis(
                analysis_id=analysis_id,
                analysis_type="process_analysis",
                evidence_id="PROC-EVIDENCE",
                tool_name="ForenseCTL Process Analyzer",
                tool_version="1.0.0",
                output_path=str(analysis_file),
                description="Análisis de procesos del sistema"
            )
            
            print("✅ Análisis de procesos completado")
            print(f"📁 Archivo: {analysis_file}")
            print(f"🆔 ID de análisis: {analysis_id}")
            print(f"📊 Total de procesos: {total_procesos}")
            print(f"⚠️ Procesos con alto CPU: {len(procesos_alto_cpu)}")
            print(f"🔴 Procesos con alta memoria: {len(procesos_alta_memoria)}")
            print(f"🛡️ Nivel de riesgo: {analisis['resumen']['nivel_riesgo']}")
            
        except Exception as e:
            print(f"❌ Error analizando procesos: {e}")
            
    def analizar_red(self):
        """Analizar conexiones de red."""
        print("\n🌐 ANALIZANDO CONEXIONES DE RED...")
        
        try:
            # Cargar evidencia de red
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/network_connections.json")
            
            if not evidence_file.exists():
                print("❌ No se encontró evidencia de red. Ejecuta la recopilación primero.")
                return
                
            with open(evidence_file, 'r', encoding='utf-8') as f:
                conexiones = json.load(f)
                
            # Análisis de red
            conexiones_establecidas = [c for c in conexiones if c['status'] == 'ESTABLISHED']
            conexiones_escuchando = [c for c in conexiones if c['status'] == 'LISTENING']
            
            # IPs externas (simulado)
            ips_externas = set()
            for conn in conexiones_establecidas:
                remote_ip = conn['remote_address'].split(':')[0]
                if not remote_ip.startswith(('127.', '192.168.', '10.', '172.')):
                    ips_externas.add(remote_ip)
                    
            analisis = {
                "total_conexiones": len(conexiones),
                "conexiones_establecidas": len(conexiones_establecidas),
                "conexiones_escuchando": len(conexiones_escuchando),
                "ips_externas": list(ips_externas),
                "timestamp": datetime.now().isoformat(),
                "resumen": {
                    "actividad_externa": len(ips_externas) > 0,
                    "puertos_abiertos": len(conexiones_escuchando),
                    "nivel_riesgo": "BAJO" if len(ips_externas) < 5 else "MEDIO"
                }
            }
            
            # Guardar análisis
            analysis_file = Path(f"cases/{self.current_case_id}/analysis/network_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analisis, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            analysis_id = f"NET-ANALYSIS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_analysis(
                analysis_id=analysis_id,
                analysis_type="network_analysis",
                evidence_id="NET-EVIDENCE",
                tool_name="ForenseCTL Network Analyzer",
                tool_version="1.0.0",
                output_path=str(analysis_file),
                description="Análisis de conexiones de red"
            )
            
            print("✅ Análisis de red completado")
            print(f"📁 Archivo: {analysis_file}")
            print(f"🆔 ID de análisis: {analysis_id}")
            print(f"🌐 Total de conexiones: {len(conexiones)}")
            print(f"✅ Conexiones establecidas: {len(conexiones_establecidas)}")
            print(f"👂 Puertos escuchando: {len(conexiones_escuchando)}")
            print(f"🌍 IPs externas: {len(ips_externas)}")
            print(f"🛡️ Nivel de riesgo: {analisis['resumen']['nivel_riesgo']}")
            
        except Exception as e:
            print(f"❌ Error analizando red: {e}")
            
    def analizar_archivos(self):
        """Analizar archivos temporales."""
        print("\n📁 ANALIZANDO ARCHIVOS...")
        
        try:
            # Cargar evidencia de archivos
            evidence_file = Path(f"cases/{self.current_case_id}/evidence/temp_files.json")
            
            if not evidence_file.exists():
                print("❌ No se encontró evidencia de archivos. Ejecuta la recopilación primero.")
                return
                
            with open(evidence_file, 'r', encoding='utf-8') as f:
                archivos = json.load(f)
                
            # Análisis de archivos
            archivos_grandes = [a for a in archivos if a['size'] > 1024*1024]  # > 1MB
            archivos_recientes = []
            
            # Archivos modificados en las últimas 24 horas
            ahora = datetime.now()
            for archivo in archivos:
                try:
                    mod_time = datetime.fromisoformat(archivo['modified'].replace('Z', '+00:00'))
                    if (ahora - mod_time.replace(tzinfo=None)).days < 1:
                        archivos_recientes.append(archivo)
                except:
                    continue
                    
            analisis = {
                "total_archivos": len(archivos),
                "archivos_grandes": archivos_grandes,
                "archivos_recientes": archivos_recientes,
                "timestamp": datetime.now().isoformat(),
                "resumen": {
                    "archivos_sospechosos": len(archivos_grandes),
                    "actividad_reciente": len(archivos_recientes),
                    "nivel_riesgo": "BAJO" if len(archivos_grandes) < 5 else "MEDIO"
                }
            }
            
            # Guardar análisis
            analysis_file = Path(f"cases/{self.current_case_id}/analysis/file_analysis.json")
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(analisis, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            analysis_id = f"FILE-ANALYSIS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_analysis(
                analysis_id=analysis_id,
                analysis_type="file_analysis",
                evidence_id="FILE-EVIDENCE",
                tool_name="ForenseCTL File Analyzer",
                tool_version="1.0.0",
                output_path=str(analysis_file),
                description="Análisis de archivos del sistema"
            )
            
            print("✅ Análisis de archivos completado")
            print(f"📁 Archivo: {analysis_file}")
            print(f"🆔 ID de análisis: {analysis_id}")
            print(f"📊 Total de archivos: {len(archivos)}")
            print(f"📈 Archivos grandes: {len(archivos_grandes)}")
            print(f"🕒 Archivos recientes: {len(archivos_recientes)}")
            print(f"🛡️ Nivel de riesgo: {analisis['resumen']['nivel_riesgo']}")
            
        except Exception as e:
            print(f"❌ Error analizando archivos: {e}")
            
    def construir_timeline(self):
        """Construir timeline de eventos."""
        print("\n⏱️ CONSTRUYENDO TIMELINE...")
        
        try:
            timeline_builder = TimelineBuilder(self.current_case_id)
            
            # Obtener eventos reales del sistema
            try:
                analyzer = RealSystemAnalyzer()
                eventos = analyzer.get_real_timeline_events()
                print(f"✅ Obtenidos {len(eventos)} eventos reales del sistema")
            except Exception as e:
                print(f"⚠️ Error obteniendo datos reales, usando datos básicos: {e}")
                eventos = [
                    {
                        "timestamp": datetime.now().isoformat(),
                        "event_type": "analysis_error",
                        "description": f"Error recopilando datos reales: {str(e)}",
                        "source": "ForenseCTL"
                    }
                ]
            
            # Guardar timeline
            timeline_file = Path(f"cases/{self.current_case_id}/analysis/timeline.json")
            with open(timeline_file, 'w', encoding='utf-8') as f:
                json.dump(eventos, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            analysis_id = f"TIMELINE-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_analysis(
                analysis_id=analysis_id,
                analysis_type="timeline",
                evidence_id="TIMELINE-EVIDENCE",
                tool_name="ForenseCTL Timeline Builder",
                tool_version="1.0.0",
                output_path=str(timeline_file),
                description="Timeline de eventos del sistema"
            )
            
            print("✅ Timeline construido")
            print(f"📁 Archivo: {timeline_file}")
            print(f"🆔 ID de análisis: {analysis_id}")
            print(f"⏱️ Total de eventos: {len(eventos)}")
            
        except Exception as e:
            print(f"❌ Error construyendo timeline: {e}")
            
    def buscar_artefactos(self):
        """Buscar artefactos específicos."""
        print("\n🔍 BUSCANDO ARTEFACTOS...")
        
        try:
            artifact_extractor = ArtifactExtractor(self.current_case_id)
            
            # Obtener artefactos reales del sistema
            try:
                analyzer = RealSystemAnalyzer()
                artefactos = analyzer.get_real_artifacts()
                print(f"✅ Artefactos extraídos del sistema real")
                print(f"  - {len(artefactos.get('registry_keys', []))} claves de registro")
                print(f"  - {len(artefactos.get('browser_artifacts', []))} artefactos de navegador")
                print(f"  - {len(artefactos.get('system_artifacts', []))} artefactos del sistema")
            except Exception as e:
                print(f"⚠️ Error obteniendo artefactos reales, usando datos básicos: {e}")
                artefactos = {
                    "error": f"Error recopilando artefactos reales: {str(e)}",
                    "timestamp": datetime.now().isoformat()
                }
            
            # Guardar artefactos
            artifacts_file = Path(f"cases/{self.current_case_id}/analysis/artifacts.json")
            with open(artifacts_file, 'w', encoding='utf-8') as f:
                json.dump(artefactos, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            analysis_id = f"ARTIFACTS-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_analysis(
                analysis_id=analysis_id,
                analysis_type="artifacts",
                evidence_id="ARTIFACTS-EVIDENCE",
                tool_name="ForenseCTL Artifact Extractor",
                tool_version="1.0.0",
                output_path=str(artifacts_file),
                description="Artefactos encontrados en el sistema"
            )
            
            print("✅ Búsqueda de artefactos completada")
            print(f"📁 Archivo: {artifacts_file}")
            print(f"🆔 ID de análisis: {analysis_id}")
            print(f"🔑 Claves de registro: {len(artefactos['registry_keys'])}")
            print(f"🌐 Artefactos de navegador: {len(artefactos['browser_artifacts'])}")
            print(f"💻 Artefactos de sistema: {len(artefactos['system_artifacts'])}")
            
        except Exception as e:
            print(f"❌ Error buscando artefactos: {e}")
            
    def analisis_seguridad(self):
        """Realizar análisis de seguridad."""
        print("\n🛡️ REALIZANDO ANÁLISIS DE SEGURIDAD...")
        
        try:
            # Realizar análisis de seguridad real del sistema
            try:
                analyzer = RealSystemAnalyzer()
                security_analysis = analyzer.get_real_security_analysis()
                print(f"✅ Análisis de seguridad completado")
                print(f"  - Nivel de riesgo: {security_analysis['risk_level']}")
                print(f"  - Puntuación de seguridad: {security_analysis['security_score']}/100")
                print(f"  - Amenazas detectadas: {len(security_analysis.get('threats_detected', []))}")
                print(f"  - Vulnerabilidades: {len(security_analysis.get('vulnerabilities', []))}")
            except Exception as e:
                print(f"⚠️ Error en análisis de seguridad real, usando datos básicos: {e}")
                security_analysis = {
                    "error": f"Error en análisis de seguridad real: {str(e)}",
                    "risk_level": "UNKNOWN",
                    "security_score": 0,
                    "timestamp": datetime.now().isoformat()
                }
            
            # Guardar análisis de seguridad
            security_file = Path(f"cases/{self.current_case_id}/analysis/security_analysis.json")
            with open(security_file, 'w', encoding='utf-8') as f:
                json.dump(security_analysis, f, indent=2, ensure_ascii=False)
                
            # Registrar en manifest
            analysis_id = f"SECURITY-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_analysis(
                analysis_id=analysis_id,
                analysis_type="security",
                evidence_id="SECURITY-EVIDENCE",
                tool_name="ForenseCTL Security Analyzer",
                tool_version="1.0.0",
                output_path=str(security_file),
                description="Análisis de seguridad del sistema"
            )
            
            print("✅ Análisis de seguridad completado")
            print(f"📁 Archivo: {security_file}")
            print(f"🆔 ID de análisis: {analysis_id}")
            print(f"🛡️ Nivel de riesgo: {security_analysis['risk_level']}")
            print(f"📊 Puntuación de seguridad: {security_analysis['security_score']}/100")
            print(f"⚠️ Amenazas detectadas: {len(security_analysis['threats_detected'])}")
            print(f"🔍 Vulnerabilidades: {len(security_analysis['vulnerabilities'])}")
            
        except Exception as e:
            print(f"❌ Error en análisis de seguridad: {e}")
            
    def analisis_completo(self):
        """Realizar análisis completo automático."""
        print("\n📈 ANÁLISIS COMPLETO AUTOMÁTICO...")
        print("Esto puede tomar unos momentos...\n")
        
        try:
            # Ejecutar todos los análisis
            self.analizar_procesos()
            print()
            self.analizar_red()
            print()
            self.analizar_archivos()
            print()
            self.construir_timeline()
            print()
            self.buscar_artefactos()
            print()
            self.analisis_seguridad()
            
            print("\n🎉 ANÁLISIS COMPLETO TERMINADO")
            print("✅ Todos los análisis han sido completados y registrados")
            
        except Exception as e:
            print(f"❌ Error en análisis completo: {e}")
            
    def generar_reportes(self):
        """Submenu para generación de reportes."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        while True:
            print(f"\n📄 GENERACIÓN DE REPORTES - {self.current_case_id}")
            print("-" * 50)
            print("1. 📋 Reporte técnico (HTML)")
            print("2. 📄 Reporte técnico (PDF)")
            print("3. 📝 Reporte técnico (DOCX)")
            print("4. 📊 Reporte ejecutivo (HTML)")
            print("5. 📈 Reporte ejecutivo (PDF)")
            print("6. 🎯 Reporte completo (todos los formatos)")
            print("7. 📋 Listar reportes generados")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.generar_reporte_html()
            elif opcion == "2":
                self.generar_reporte_pdf()
            elif opcion == "3":
                self.generar_reporte_docx()
            elif opcion == "4":
                self.generar_reporte_ejecutivo_html()
            elif opcion == "5":
                self.generar_reporte_ejecutivo_pdf()
            elif opcion == "6":
                self.generar_reportes_completos()
            elif opcion == "7":
                self.listar_reportes()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def generar_reporte_html(self):
        """Generar reporte técnico en HTML."""
        print("\n📋 GENERANDO REPORTE TÉCNICO HTML...")
        
        try:
            report_generator = ReportGenerator(self.current_case_id, examiner="Analista Forense")
            
            report_info = report_generator.generate_report(
                report_type="technical",
                output_format="html",
                template_name="technical_report_es",
                language="es"
            )
            
            print("✅ Reporte HTML generado")
            print(f"📁 Archivo: {report_info['output_file']}")
            print(f"🆔 ID de reporte: {report_info['report_id']}")
            print(f"📊 Formato: {report_info['format']}")
            
        except Exception as e:
            print(f"❌ Error generando reporte HTML: {e}")
    
    def _get_analysis_data(self):
        """Recopila datos de análisis desde los archivos JSON"""
        analysis_dir = Path(f"cases/{self.current_case_id}/analysis")
        data = {
            'analyses': [],
            'timeline_events': [],
            'artifacts': {},
            'security': {}
        }
        
        try:
            # Leer timeline
            timeline_file = analysis_dir / "timeline.json"
            if timeline_file.exists():
                with open(timeline_file, 'r', encoding='utf-8') as f:
                    timeline_data = json.load(f)
                    # timeline.json contiene un array directamente
                    data['timeline_events'] = timeline_data if isinstance(timeline_data, list) else []
                    data['analyses'].append({
                        'type': 'Timeline Construction',
                        'description': f"Análisis temporal con {len(data['timeline_events'])} eventos"
                    })
            
            # Leer artefactos
            artifacts_file = analysis_dir / "artifacts.json"
            if artifacts_file.exists():
                with open(artifacts_file, 'r', encoding='utf-8') as f:
                    artifacts_data = json.load(f)
                    # artifacts.json contiene un objeto con las claves directamente
                    data['artifacts'] = artifacts_data
                    data['analyses'].append({
                        'type': 'Artifact Search',
                        'description': f"Búsqueda de artefactos del sistema"
                    })
            
            # Leer análisis de seguridad
            security_file = analysis_dir / "security_analysis.json"
            if security_file.exists():
                with open(security_file, 'r', encoding='utf-8') as f:
                    security_data = json.load(f)
                    # security_analysis.json contiene un objeto con los datos directamente
                    data['security'] = security_data
                    data['analyses'].append({
                        'type': 'Security Analysis',
                        'description': f"Análisis de seguridad del sistema"
                    })
                    
        except Exception as e:
            print(f"Error leyendo datos de análisis: {e}")
            
        return data
            
    def generar_reporte_pdf(self):
        """Generar reporte técnico en PDF."""
        print("\n📄 GENERANDO REPORTE TÉCNICO PDF...")
        
        try:
            # Usar el sistema de exportación alternativo
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Crear reporte PDF
            reports_dir = Path(f"cases/{self.current_case_id}/reports")
            reports_dir.mkdir(parents=True, exist_ok=True)
            
            pdf_file = reports_dir / f"{self.current_case_id}_reporte_tecnico.pdf"
            
            doc = SimpleDocTemplate(str(pdf_file), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Obtener datos de análisis reales
            analysis_data = self._get_analysis_data()
            
            # Contenido del reporte
            story.append(Paragraph("Reporte Técnico de Análisis Forense", styles['Title']))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"Caso: {self.current_case_id}", styles['Heading2']))
            story.append(Paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Resumen Ejecutivo
            story.append(Paragraph("Resumen Ejecutivo", styles['Heading2']))
            story.append(Paragraph("Este reporte presenta los hallazgos del análisis forense digital realizado.", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Análisis realizados
            if analysis_data['analyses']:
                story.append(Paragraph("Análisis Realizados", styles['Heading2']))
                for analysis in analysis_data['analyses']:
                    story.append(Paragraph(f"• {analysis['type']}: {analysis['description']}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Timeline
            if analysis_data['timeline_events']:
                story.append(Paragraph("Timeline de Eventos", styles['Heading2']))
                story.append(Paragraph(f"Total de eventos analizados: {len(analysis_data['timeline_events'])}", styles['Normal']))
                for event in analysis_data['timeline_events'][:5]:  # Mostrar primeros 5
                    story.append(Paragraph(f"• {event['timestamp']}: {event['description']}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Artefactos
            if analysis_data['artifacts']:
                story.append(Paragraph("Artefactos Encontrados", styles['Heading2']))
                artifacts = analysis_data['artifacts']
                story.append(Paragraph(f"• Claves de registro: {len(artifacts.get('registry_keys', []))}", styles['Normal']))
                story.append(Paragraph(f"• Artefactos de navegador: {len(artifacts.get('browser_artifacts', []))}", styles['Normal']))
                story.append(Paragraph(f"• Artefactos de sistema: {len(artifacts.get('system_artifacts', []))}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Análisis de seguridad
            if analysis_data['security']:
                story.append(Paragraph("Análisis de Seguridad", styles['Heading2']))
                security = analysis_data['security']
                story.append(Paragraph(f"• Nivel de riesgo: {security.get('risk_level', 'N/A')}", styles['Normal']))
                story.append(Paragraph(f"• Puntuación de seguridad: {security.get('security_score', 'N/A')}/100", styles['Normal']))
                story.append(Paragraph(f"• Amenazas detectadas: {len(security.get('threats_detected', []))}", styles['Normal']))
                story.append(Spacer(1, 12))
            
            doc.build(story)
            
            # Registrar en manifest
            report_id = f"REPORT-PDF-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_report(
                report_id=report_id,
                report_type="technical",
                report_format="pdf",
                output_path=str(pdf_file),
                examiner="Analista Forense",
                description="Reporte técnico en formato PDF"
            )
            
            print("✅ Reporte PDF generado")
            print(f"📁 Archivo: {pdf_file}")
            print(f"🆔 ID de reporte: {report_id}")
            
        except Exception as e:
            print(f"❌ Error generando reporte PDF: {e}")
            
    def generar_reporte_docx(self):
        """Generar reporte técnico en DOCX."""
        print("\n📝 GENERANDO REPORTE TÉCNICO DOCX...")
        
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Obtener datos de análisis reales
            analysis_data = self._get_analysis_data()
            
            # Crear documento
            doc = Document()
            
            # Título
            title = doc.add_heading('Reporte Técnico de Análisis Forense', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del caso
            doc.add_heading('Información del Caso', level=1)
            p = doc.add_paragraph()
            p.add_run('Caso: ').bold = True
            p.add_run(f"{self.current_case_id}\n")
            p.add_run('Generado: ').bold = True
            p.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
            p.add_run('Examinador: ').bold = True
            p.add_run('Analista Forense')
            
            # Resumen ejecutivo
            doc.add_heading('Resumen Ejecutivo', level=1)
            doc.add_paragraph(
                'Este reporte presenta los hallazgos del análisis forense digital realizado '
                'sobre el sistema objetivo. Se han examinado procesos, conexiones de red, '
                'archivos y otros artefactos para proporcionar una evaluación completa.'
            )
            
            # Análisis realizados
            if analysis_data['analyses']:
                doc.add_heading('Análisis Realizados', level=1)
                for analysis in analysis_data['analyses']:
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(f"{analysis['type']}: {analysis['description']}")
            
            # Timeline
            if analysis_data['timeline_events']:
                doc.add_heading('Timeline de Eventos', level=1)
                doc.add_paragraph(f"Total de eventos analizados: {len(analysis_data['timeline_events'])}")
                for event in analysis_data['timeline_events'][:5]:  # Mostrar primeros 5
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(f"{event['timestamp']}: {event['description']}")
            
            # Artefactos
            if analysis_data['artifacts']:
                doc.add_heading('Artefactos Encontrados', level=1)
                artifacts = analysis_data['artifacts']
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(f"Claves de registro: {len(artifacts.get('registry_keys', []))}\n")
                p.add_run('• ').bold = True
                p.add_run(f"Artefactos de navegador: {len(artifacts.get('browser_artifacts', []))}\n")
                p.add_run('• ').bold = True
                p.add_run(f"Artefactos de sistema: {len(artifacts.get('system_artifacts', []))}")
            
            # Análisis de seguridad
            if analysis_data['security']:
                doc.add_heading('Análisis de Seguridad', level=1)
                security = analysis_data['security']
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(f"Nivel de riesgo: {security.get('risk_level', 'N/A')}\n")
                p.add_run('• ').bold = True
                p.add_run(f"Puntuación de seguridad: {security.get('security_score', 'N/A')}/100\n")
                p.add_run('• ').bold = True
                p.add_run(f"Amenazas detectadas: {len(security.get('threats_detected', []))}")
            
            # Guardar documento
            reports_dir = Path(f"cases/{self.current_case_id}/reports")
            reports_dir.mkdir(parents=True, exist_ok=True)
            
            docx_file = reports_dir / f"{self.current_case_id}_reporte_tecnico.docx"
            doc.save(str(docx_file))
            
            # Registrar en manifest
            report_id = f"REPORT-DOCX-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_report(
                report_id=report_id,
                report_type="technical",
                report_format="docx",
                output_path=str(docx_file),
                examiner="Analista Forense",
                description="Reporte técnico en formato DOCX"
            )
            
            print("✅ Reporte DOCX generado")
            print(f"📁 Archivo: {docx_file}")
            print(f"🆔 ID de reporte: {report_id}")
            
        except Exception as e:
            print(f"❌ Error generando reporte DOCX: {e}")
            
    def generar_reporte_ejecutivo_html(self):
        """Generar reporte ejecutivo en HTML."""
        print("\n📊 GENERANDO REPORTE EJECUTIVO HTML...")
        
        try:
            report_generator = ReportGenerator(self.current_case_id, examiner="Analista Forense")
            
            report_info = report_generator.generate_report(
                report_type="executive",
                output_format="html",
                template_name="executive_report_es",
                language="es"
            )
            
            print("✅ Reporte ejecutivo HTML generado")
            print(f"📁 Archivo: {report_info['output_file']}")
            print(f"🆔 ID de reporte: {report_info['report_id']}")
            
        except Exception as e:
            print(f"❌ Error generando reporte ejecutivo HTML: {e}")
            
    def generar_reporte_ejecutivo_pdf(self):
        """Generar reporte ejecutivo en PDF."""
        print("\n📈 GENERANDO REPORTE EJECUTIVO PDF...")
        
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Crear reporte PDF ejecutivo
            reports_dir = Path(f"cases/{self.current_case_id}/reports")
            reports_dir.mkdir(parents=True, exist_ok=True)
            
            pdf_file = reports_dir / f"{self.current_case_id}_reporte_ejecutivo.pdf"
            
            doc = SimpleDocTemplate(str(pdf_file), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Obtener datos de análisis reales
            analysis_data = self._get_analysis_data()
            
            # Contenido ejecutivo
            story.append(Paragraph("Reporte Ejecutivo - Análisis Forense", styles['Title']))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"Caso: {self.current_case_id}", styles['Heading2']))
            story.append(Paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Resumen para la Dirección
            story.append(Paragraph("Resumen para la Dirección", styles['Heading2']))
            story.append(Paragraph("El análisis forense digital ha sido completado exitosamente. "
                                 "Este reporte presenta los hallazgos principales y recomendaciones "
                                 "para la toma de decisiones estratégicas.", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Estado del análisis
            story.append(Paragraph("Estado del Análisis", styles['Heading2']))
            total_analyses = len(analysis_data['analyses'])
            story.append(Paragraph(f"• Análisis completados: {total_analyses}", styles['Normal']))
            
            if analysis_data['timeline_events']:
                story.append(Paragraph(f"• Eventos analizados: {len(analysis_data['timeline_events'])}", styles['Normal']))
            
            if analysis_data['security']:
                security = analysis_data['security']
                risk_level = security.get('risk_level', 'BAJO')
                story.append(Paragraph(f"• Nivel de riesgo identificado: {risk_level}", styles['Normal']))
                
                if security.get('security_score'):
                    story.append(Paragraph(f"• Puntuación de seguridad: {security['security_score']}/100", styles['Normal']))
            
            story.append(Spacer(1, 12))
            
            # Hallazgos principales
            story.append(Paragraph("Hallazgos Principales", styles['Heading2']))
            if analysis_data['artifacts']:
                artifacts = analysis_data['artifacts']
                total_artifacts = (len(artifacts.get('registry_keys', [])) + 
                                 len(artifacts.get('browser_artifacts', [])) + 
                                 len(artifacts.get('system_artifacts', [])))
                story.append(Paragraph(f"• Total de artefactos identificados: {total_artifacts}", styles['Normal']))
            
            if analysis_data['security'] and analysis_data['security'].get('threats_detected'):
                threats = len(analysis_data['security']['threats_detected'])
                story.append(Paragraph(f"• Amenazas potenciales detectadas: {threats}", styles['Normal']))
            else:
                story.append(Paragraph("• No se detectaron amenazas críticas inmediatas", styles['Normal']))
            
            story.append(Spacer(1, 12))
            
            # Recomendaciones
            story.append(Paragraph("Recomendaciones", styles['Heading2']))
            story.append(Paragraph("• Continuar con el monitoreo regular del sistema", styles['Normal']))
            story.append(Paragraph("• Mantener las medidas de seguridad actuales", styles['Normal']))
            story.append(Paragraph("• Revisar y actualizar las políticas de seguridad", styles['Normal']))
            
            if analysis_data['security'] and analysis_data['security'].get('risk_level') in ['ALTO', 'CRÍTICO']:
                story.append(Paragraph("• ACCIÓN INMEDIATA: Implementar medidas de seguridad adicionales", styles['Normal']))
                story.append(Paragraph("• Realizar análisis forense más profundo", styles['Normal']))
            
            doc.build(story)
            
            # Registrar en manifest
            report_id = f"EXEC-PDF-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.manifest.register_report(
                report_id=report_id,
                report_type="executive",
                report_format="pdf",
                output_path=str(pdf_file),
                examiner="Analista Forense",
                description="Reporte ejecutivo en formato PDF"
            )
            
            print("✅ Reporte ejecutivo PDF generado")
            print(f"📁 Archivo: {pdf_file}")
            print(f"🆔 ID de reporte: {report_id}")
            
        except Exception as e:
            print(f"❌ Error generando reporte ejecutivo PDF: {e}")
            
    def generar_reportes_completos(self):
        """Generar reportes en todos los formatos."""
        print("\n🎯 GENERANDO REPORTES COMPLETOS...")
        print("Esto puede tomar unos momentos...\n")
        
        try:
            # Generar todos los formatos
            self.generar_reporte_html()
            print()
            self.generar_reporte_pdf()
            print()
            self.generar_reporte_docx()
            print()
            self.generar_reporte_ejecutivo_html()
            print()
            self.generar_reporte_ejecutivo_pdf()
            
            print("\n🎉 REPORTES COMPLETOS GENERADOS")
            print("✅ Todos los reportes han sido generados exitosamente")
            
        except Exception as e:
            print(f"❌ Error generando reportes completos: {e}")
            
    def listar_reportes(self):
        """Listar todos los reportes generados."""
        print("\n📋 REPORTES GENERADOS")
        print("-" * 30)
        
        try:
            reports_dir = Path(f"cases/{self.current_case_id}/reports")
            
            if not reports_dir.exists():
                print("📭 No hay reportes generados")
                return
                
            reportes = list(reports_dir.glob("*"))
            
            if not reportes:
                print("📭 No hay reportes en el directorio")
                return
                
            for i, reporte in enumerate(reportes, 1):
                size = reporte.stat().st_size if reporte.is_file() else 0
                modified = datetime.fromtimestamp(reporte.stat().st_mtime).strftime('%d/%m/%Y %H:%M')
                
                print(f"{i}. 📄 {reporte.name}")
                print(f"   📊 Tamaño: {size:,} bytes")
                print(f"   📅 Modificado: {modified}")
                print()
                
        except Exception as e:
            print(f"❌ Error listando reportes: {e}")
            
    def cadena_custodia(self):
        """Submenu para cadena de custodia."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        while True:
            print(f"\n🔗 CADENA DE CUSTODIA - {self.current_case_id}")
            print("-" * 45)
            print("1. 📝 Registrar nueva entrada")
            print("2. 📋 Ver historial completo")
            print("3. 🔍 Buscar por evidencia")
            print("4. 📊 Estadísticas de custodia")
            print("5. 📄 Exportar cadena de custodia")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.registrar_entrada_custodia()
            elif opcion == "2":
                self.ver_historial_custodia()
            elif opcion == "3":
                self.buscar_evidencia_custodia()
            elif opcion == "4":
                self.estadisticas_custodia()
            elif opcion == "5":
                self.exportar_custodia()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def registrar_entrada_custodia(self):
        """Registrar nueva entrada en cadena de custodia."""
        print("\n📝 REGISTRAR ENTRADA DE CUSTODIA")
        print("-" * 40)
        
        evidence_id = input("🆔 ID de evidencia: ").strip()
        action = input("⚙️ Acción realizada: ").strip()
        handler = input("👤 Responsable: ").strip() or "Analista Forense"
        notes = input("📝 Notas adicionales: ").strip()
        
        if not evidence_id or not action:
            print("❌ ID de evidencia y acción son requeridos")
            return
            
        try:
            entry_id = self.chain_of_custody.add_entry(
                evidence_id=evidence_id,
                action=action,
                handler=handler,
                notes=notes
            )
            
            print("✅ Entrada registrada en cadena de custodia")
            print(f"🆔 ID de entrada: {entry_id}")
            print(f"📋 Evidencia: {evidence_id}")
            print(f"⚙️ Acción: {action}")
            print(f"👤 Responsable: {handler}")
            
        except Exception as e:
            print(f"❌ Error registrando entrada: {e}")
            
    def ver_historial_custodia(self):
        """Ver historial completo de cadena de custodia."""
        print("\n📋 HISTORIAL DE CADENA DE CUSTODIA")
        print("-" * 45)
        
        try:
            historial = self.chain_of_custody.get_chain_history()
            
            if not historial:
                print("📭 No hay entradas en la cadena de custodia")
                return
                
            for i, entrada in enumerate(historial, 1):
                print(f"{i}. 🆔 {entrada.get('entry_id', 'N/A')}")
                print(f"   📋 Evidencia: {entrada.get('evidence_id', 'N/A')}")
                print(f"   ⚙️ Acción: {entrada.get('action', 'N/A')}")
                print(f"   👤 Responsable: {entrada.get('handler', 'N/A')}")
                print(f"   📅 Fecha: {entrada.get('timestamp', 'N/A')[:19]}")
                if entrada.get('notes'):
                    print(f"   📝 Notas: {entrada['notes']}")
                print()
                
        except Exception as e:
            print(f"❌ Error obteniendo historial: {e}")
            
    def buscar_evidencia_custodia(self):
        """Buscar entradas por ID de evidencia."""
        print("\n🔍 BUSCAR EN CADENA DE CUSTODIA")
        print("-" * 35)
        
        evidence_id = input("🆔 ID de evidencia a buscar: ").strip()
        
        if not evidence_id:
            print("❌ ID de evidencia requerido")
            return
            
        try:
            entradas = self.chain_of_custody.get_evidence_history(evidence_id)
            
            if not entradas:
                print(f"📭 No se encontraron entradas para {evidence_id}")
                return
                
            print(f"\n📋 HISTORIAL DE {evidence_id}")
            print("-" * 30)
            
            for i, entrada in enumerate(entradas, 1):
                print(f"{i}. ⚙️ {entrada.get('action', 'N/A')}")
                print(f"   👤 {entrada.get('handler', 'N/A')}")
                print(f"   📅 {entrada.get('timestamp', 'N/A')[:19]}")
                if entrada.get('notes'):
                    print(f"   📝 {entrada['notes']}")
                print()
                
        except Exception as e:
            print(f"❌ Error buscando evidencia: {e}")
            
    def estadisticas_custodia(self):
        """Mostrar estadísticas de cadena de custodia."""
        print("\n📊 ESTADÍSTICAS DE CADENA DE CUSTODIA")
        print("-" * 45)
        
        try:
            historial = self.chain_of_custody.get_chain_history()
            
            if not historial:
                print("📭 No hay datos para mostrar estadísticas")
                return
                
            # Estadísticas básicas
            total_entradas = len(historial)
            evidencias_unicas = len(set(e.get('evidence_id', '') for e in historial))
            responsables_unicos = len(set(e.get('handler', '') for e in historial))
            
            # Acciones más comunes
            acciones = {}
            for entrada in historial:
                accion = entrada.get('action', 'N/A')
                acciones[accion] = acciones.get(accion, 0) + 1
                
            print(f"📊 Total de entradas: {total_entradas}")
            print(f"🆔 Evidencias únicas: {evidencias_unicas}")
            print(f"👤 Responsables únicos: {responsables_unicos}")
            print("\n⚙️ ACCIONES MÁS COMUNES:")
            
            for accion, count in sorted(acciones.items(), key=lambda x: x[1], reverse=True):
                print(f"   {accion}: {count} veces")
                
        except Exception as e:
            print(f"❌ Error calculando estadísticas: {e}")
            
    def exportar_custodia(self):
        """Exportar cadena de custodia."""
        print("\n📄 EXPORTAR CADENA DE CUSTODIA")
        print("-" * 35)
        
        try:
            historial = self.chain_of_custody.get_chain_history()
            
            if not historial:
                print("📭 No hay datos para exportar")
                return
                
            # Exportar a JSON
            exports_dir = Path(f"cases/{self.current_case_id}/exports")
            exports_dir.mkdir(parents=True, exist_ok=True)
            
            export_file = exports_dir / f"cadena_custodia_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            
            with open(export_file, 'w', encoding='utf-8') as f:
                json.dump(historial, f, indent=2, ensure_ascii=False)
                
            print("✅ Cadena de custodia exportada")
            print(f"📁 Archivo: {export_file}")
            print(f"📊 Total de entradas: {len(historial)}")
            
        except Exception as e:
            print(f"❌ Error exportando cadena de custodia: {e}")
            
    def exportacion_datos(self):
        """Submenu para exportación de datos."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        while True:
            print(f"\n📊 EXPORTACIÓN DE DATOS - {self.current_case_id}")
            print("-" * 45)
            print("1. 📋 Exportar evidencia")
            print("2. 📊 Exportar análisis")
            print("3. 📄 Exportar reportes")
            print("4. 🔗 Exportar cadena de custodia")
            print("5. 📦 Exportación completa")
            print("6. 📋 Listar exportaciones")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.exportar_evidencia()
            elif opcion == "2":
                self.exportar_analisis()
            elif opcion == "3":
                self.exportar_reportes_datos()
            elif opcion == "4":
                self.exportar_custodia()
            elif opcion == "5":
                self.exportacion_completa()
            elif opcion == "6":
                self.listar_exportaciones()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def exportar_evidencia(self):
        """Exportar toda la evidencia recopilada."""
        print("\n📋 EXPORTANDO EVIDENCIA...")
        
        try:
            export_manager = ExportManager(self.current_case_id)
            
            export_info = export_manager.export_evidence(
                format="json",
                include_metadata=True
            )
            
            print("✅ Evidencia exportada")
            print(f"📁 Archivo: {export_info['output_file']}")
            print(f"📊 Total de archivos: {export_info.get('file_count', 0)}")
            
        except Exception as e:
            print(f"❌ Error exportando evidencia: {e}")
            
    def exportar_analisis(self):
        """Exportar todos los análisis realizados."""
        print("\n📊 EXPORTANDO ANÁLISIS...")
        
        try:
            export_manager = ExportManager(self.current_case_id)
            
            export_info = export_manager.export_analysis(
                format="json",
                include_metadata=True
            )
            
            print("✅ Análisis exportados")
            print(f"📁 Archivo: {export_info['output_file']}")
            print(f"📊 Total de análisis: {export_info.get('analysis_count', 0)}")
            
        except Exception as e:
            print(f"❌ Error exportando análisis: {e}")
            
    def exportar_reportes_datos(self):
        """Exportar metadatos de reportes."""
        print("\n📄 EXPORTANDO DATOS DE REPORTES...")
        
        try:
            export_manager = ExportManager(self.current_case_id)
            
            export_info = export_manager.export_reports(
                format="json",
                include_content=False
            )
            
            print("✅ Datos de reportes exportados")
            print(f"📁 Archivo: {export_info['output_file']}")
            print(f"📊 Total de reportes: {export_info.get('report_count', 0)}")
            
        except Exception as e:
            print(f"❌ Error exportando datos de reportes: {e}")
            
    def exportacion_completa(self):
        """Realizar exportación completa del caso."""
        print("\n📦 EXPORTACIÓN COMPLETA...")
        print("Esto puede tomar unos momentos...\n")
        
        try:
            export_manager = ExportManager(self.current_case_id)
            
            # Exportar todo
            self.exportar_evidencia()
            print()
            self.exportar_analisis()
            print()
            self.exportar_reportes_datos()
            print()
            self.exportar_custodia()
            
            print("\n🎉 EXPORTACIÓN COMPLETA TERMINADA")
            print("✅ Todos los datos han sido exportados")
            
        except Exception as e:
            print(f"❌ Error en exportación completa: {e}")
            
    def listar_exportaciones(self):
        """Listar todas las exportaciones realizadas."""
        print("\n📋 EXPORTACIONES REALIZADAS")
        print("-" * 35)
        
        try:
            exports_dir = Path(f"cases/{self.current_case_id}/exports")
            
            if not exports_dir.exists():
                print("📭 No hay exportaciones realizadas")
                return
                
            exportaciones = list(exports_dir.glob("*"))
            
            if not exportaciones:
                print("📭 No hay archivos de exportación")
                return
                
            for i, export_file in enumerate(exportaciones, 1):
                size = export_file.stat().st_size if export_file.is_file() else 0
                modified = datetime.fromtimestamp(export_file.stat().st_mtime).strftime('%d/%m/%Y %H:%M')
                
                print(f"{i}. 📄 {export_file.name}")
                print(f"   📊 Tamaño: {size:,} bytes")
                print(f"   📅 Creado: {modified}")
                print()
                
        except Exception as e:
            print(f"❌ Error listando exportaciones: {e}")
            
    def estado_sistema(self):
        """Mostrar estado general del sistema."""
        print("\n📈 ESTADO DEL SISTEMA FORENSECTL")
        print("=" * 50)
        
        try:
            # Estado del caso actual
            if self.current_case_id:
                print(f"📁 Caso activo: {self.current_case_id}")
                
                # Contar archivos en cada directorio
                case_dir = Path(f"cases/{self.current_case_id}")
                
                evidence_count = len(list((case_dir / "evidence").glob("*"))) if (case_dir / "evidence").exists() else 0
                analysis_count = len(list((case_dir / "analysis").glob("*"))) if (case_dir / "analysis").exists() else 0
                reports_count = len(list((case_dir / "reports").glob("*"))) if (case_dir / "reports").exists() else 0
                exports_count = len(list((case_dir / "exports").glob("*"))) if (case_dir / "exports").exists() else 0
                
                print(f"🔍 Archivos de evidencia: {evidence_count}")
                print(f"📊 Archivos de análisis: {analysis_count}")
                print(f"📄 Reportes generados: {reports_count}")
                print(f"📦 Exportaciones: {exports_count}")
                
                # Estado de la cadena de custodia
                if self.chain_of_custody:
                    historial = self.chain_of_custody.get_chain_history()
                    print(f"🔗 Entradas en cadena de custodia: {len(historial)}")
                    
            else:
                print("❌ No hay caso activo seleccionado")
                
            # Estado general
            print("\n🛠️ ESTADO DE COMPONENTES:")
            print(f"✅ Case Manager: {'Activo' if self.case_manager else 'Inactivo'}")
            print(f"✅ Chain of Custody: {'Activo' if self.chain_of_custody else 'Inactivo'}")
            print(f"✅ Manifest: {'Activo' if self.manifest else 'Inactivo'}")
            
            # Información del sistema
            print("\n💻 INFORMACIÓN DEL SISTEMA:")
            print(f"🖥️ Plataforma: {sys.platform}")
            print(f"🐍 Python: {sys.version.split()[0]}")
            print(f"📂 Directorio de trabajo: {Path.cwd()}")
            
        except Exception as e:
            print(f"❌ Error obteniendo estado del sistema: {e}")
            
    def herramientas_adicionales(self):
        """Submenu para herramientas adicionales."""
        while True:
            print("\n🛠️ HERRAMIENTAS ADICIONALES")
            print("-" * 35)
            print("1. 🧹 Limpiar archivos temporales")
            print("2. 🔍 Verificar integridad de archivos")
            print("3. 📊 Estadísticas del caso")
            print("4. 🔧 Configuración del sistema")
            print("5. 📋 Información de dependencias")
            print("6. 🗑️ Herramientas de limpieza")
            print("0. ⬅️ Volver al menú principal")
            
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.limpiar_temporales()
            elif opcion == "2":
                self.verificar_integridad()
            elif opcion == "3":
                self.estadisticas_caso()
            elif opcion == "4":
                self.configuracion_sistema()
            elif opcion == "5":
                self.info_dependencias()
            elif opcion == "6":
                self.herramientas_limpieza()
            elif opcion == "0":
                break
            else:
                print("❌ Opción inválida")
                
    def limpiar_temporales(self):
        """Limpiar archivos temporales del caso."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        print("\n🧹 LIMPIANDO ARCHIVOS TEMPORALES...")
        
        try:
            temp_dir = Path(f"cases/{self.current_case_id}/temp")
            
            if not temp_dir.exists():
                print("✅ No hay archivos temporales para limpiar")
                return
                
            archivos_eliminados = 0
            for archivo in temp_dir.glob("*"):
                if archivo.is_file():
                    archivo.unlink()
                    archivos_eliminados += 1
                    
            print(f"✅ Limpieza completada")
            print(f"🗑️ Archivos eliminados: {archivos_eliminados}")
            
        except Exception as e:
            print(f"❌ Error limpiando temporales: {e}")
            
    def verificar_integridad(self):
        """Verificar integridad de archivos del caso."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        print("\n🔍 VERIFICANDO INTEGRIDAD DE ARCHIVOS...")
        
        try:
            case_dir = Path(f"cases/{self.current_case_id}")
            archivos_verificados = 0
            archivos_corruptos = 0
            
            for archivo in case_dir.rglob("*"):
                if archivo.is_file() and archivo.suffix in ['.json', '.txt', '.html']:
                    try:
                        with open(archivo, 'r', encoding='utf-8') as f:
                            f.read()
                        archivos_verificados += 1
                    except Exception:
                        archivos_corruptos += 1
                        
            print(f"✅ Verificación completada")
            print(f"📊 Archivos verificados: {archivos_verificados}")
            print(f"❌ Archivos corruptos: {archivos_corruptos}")
            
            if archivos_corruptos == 0:
                print("🎉 Todos los archivos están íntegros")
            else:
                print("⚠️ Se encontraron archivos corruptos")
                
        except Exception as e:
            print(f"❌ Error verificando integridad: {e}")
            
    def estadisticas_caso(self):
        """Mostrar estadísticas detalladas del caso."""
        if not self.current_case_id:
            print("❌ Selecciona un caso activo primero")
            return
            
        print(f"\n📊 ESTADÍSTICAS DEL CASO: {self.current_case_id}")
        print("-" * 60)
        
        try:
            case_dir = Path(f"cases/{self.current_case_id}")
            
            # Estadísticas de archivos
            total_archivos = 0
            total_tamaño = 0
            
            for archivo in case_dir.rglob("*"):
                if archivo.is_file():
                    total_archivos += 1
                    total_tamaño += archivo.stat().st_size
                    
            # Estadísticas por directorio
            directorios = ['evidence', 'analysis', 'reports', 'exports', 'temp']
            
            print(f"📁 Total de archivos: {total_archivos}")
            print(f"💾 Tamaño total: {total_tamaño:,} bytes ({total_tamaño/1024/1024:.2f} MB)")
            print("\n📂 ARCHIVOS POR DIRECTORIO:")
            
            for directorio in directorios:
                dir_path = case_dir / directorio
                if dir_path.exists():
                    archivos_dir = len(list(dir_path.glob("*")))
                    print(f"   {directorio}: {archivos_dir} archivos")
                else:
                    print(f"   {directorio}: 0 archivos")
                    
            # Información temporal
            case_info = self.case_manager.get_case_info(self.current_case_id)
            created_at = case_info.get('created_at', '')
            
            if created_at:
                try:
                    created_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                    tiempo_transcurrido = datetime.now() - created_date.replace(tzinfo=None)
                    print(f"\n⏱️ Tiempo transcurrido: {tiempo_transcurrido.days} días, {tiempo_transcurrido.seconds//3600} horas")
                except:
                    pass
                    
        except Exception as e:
            print(f"❌ Error calculando estadísticas: {e}")
            
    def configuracion_sistema(self):
        """Mostrar configuración del sistema."""
        print("\n🔧 CONFIGURACIÓN DEL SISTEMA")
        print("-" * 40)
        
        try:
            print("🐍 PYTHON:")
            print(f"   Versión: {sys.version}")
            print(f"   Ejecutable: {sys.executable}")
            print(f"   Plataforma: {sys.platform}")
            
            print("\n📂 DIRECTORIOS:")
            print(f"   Trabajo: {Path.cwd()}")
            print(f"   Casos: {Path.cwd() / 'cases'}")
            print(f"   Templates: {Path.cwd() / 'forensectl' / 'templates'}")
            
            print("\n🌍 VARIABLES DE ENTORNO:")
            env_vars = ['PATH', 'PYTHONPATH', 'TEMP', 'USERNAME']
            for var in env_vars:
                value = os.environ.get(var, 'No definida')
                if len(value) > 50:
                    value = value[:50] + "..."
                print(f"   {var}: {value}")
                
        except Exception as e:
            print(f"❌ Error obteniendo configuración: {e}")
            
    def info_dependencias(self):
        """Mostrar información de dependencias."""
        print("\n📋 INFORMACIÓN DE DEPENDENCIAS")
        print("-" * 40)
        
        dependencias = {
            'jinja2': 'Templates de reportes',
            'reportlab': 'Generación de PDF',
            'python-docx': 'Generación de DOCX',
            'beautifulsoup4': 'Procesamiento HTML',
            'markdown': 'Procesamiento Markdown'
        }
        
        print("📦 DEPENDENCIAS PRINCIPALES:")
        for dep, desc in dependencias.items():
            try:
                __import__(dep.replace('-', '_'))
                status = "✅ Instalado"
            except ImportError:
                status = "❌ No instalado"
            print(f"   {dep}: {desc} - {status}")
            
        print("\n💡 DEPENDENCIAS OPCIONALES:")
        opcionales = {
            'psutil': 'Información detallada del sistema',
            'yara-python': 'Análisis con reglas YARA',
            'volatility3': 'Análisis de memoria'
        }
        
        for dep, desc in opcionales.items():
            try:
                __import__(dep.replace('-', '_'))
                status = "✅ Instalado"
            except ImportError:
                status = "❌ No instalado"
            print(f"   {dep}: {desc} - {status}")
            
    def herramientas_limpieza(self):
        """Herramientas de limpieza del sistema."""
        print("\n🗑️ HERRAMIENTAS DE LIMPIEZA")
        print("-" * 35)
        
        print("1. 🧹 Limpiar todos los archivos temporales")
        print("2. 🗑️ Eliminar casos antiguos")
        print("3. 📊 Optimizar base de datos")
        print("4. 🔄 Reiniciar configuración")
        print("0. ⬅️ Volver")
        
        opcion = input("\nSelecciona una opción: ").strip()
        
        if opcion == "1":
            self.limpiar_todos_temporales()
        elif opcion == "2":
            self.eliminar_casos_antiguos()
        elif opcion == "3":
            print("✅ Base de datos optimizada (simulado)")
        elif opcion == "4":
            print("✅ Configuración reiniciada (simulado)")
            
    def limpiar_todos_temporales(self):
        """Limpiar archivos temporales de todos los casos."""
        print("\n🧹 LIMPIANDO TODOS LOS TEMPORALES...")
        
        try:
            cases_dir = Path("cases")
            archivos_eliminados = 0
            
            for case_dir in cases_dir.glob("CASE-*"):
                temp_dir = case_dir / "temp"
                if temp_dir.exists():
                    for archivo in temp_dir.glob("*"):
                        if archivo.is_file():
                            archivo.unlink()
                            archivos_eliminados += 1
                            
            print(f"✅ Limpieza global completada")
            print(f"🗑️ Total de archivos eliminados: {archivos_eliminados}")
            
        except Exception as e:
            print(f"❌ Error en limpieza global: {e}")
            
    def eliminar_casos_antiguos(self):
        """Eliminar casos antiguos."""
        print("\n🗑️ ELIMINAR CASOS ANTIGUOS")
        print("-" * 30)
        
        dias = input("📅 Eliminar casos más antiguos que (días): ").strip()
        
        try:
            dias = int(dias)
            if dias <= 0:
                print("❌ Número de días debe ser positivo")
                return
                
            confirmacion = input(f"⚠️ ¿Eliminar casos más antiguos que {dias} días? (sí/no): ").strip().lower()
            
            if confirmacion in ['sí', 'si', 'yes', 'y']:
                print(f"✅ Casos antiguos eliminados (simulado)")
                print(f"📊 Se eliminarían casos más antiguos que {dias} días")
            else:
                print("❌ Eliminación cancelada")
                
        except ValueError:
            print("❌ Número de días inválido")
            
    def analisis_forense_completo(self):
        """Realizar análisis forense completo del sistema actual."""
        print("\n🔬 ANÁLISIS FORENSE COMPLETO DEL SISTEMA")
        print("=" * 50)
        
        try:
            print("🔍 Iniciando análisis forense completo...")
            analyzer = RealSystemAnalyzer()
            
            print("📊 Recopilando datos del sistema...")
            comprehensive_data = analyzer.get_comprehensive_analysis()
            
            # Crear directorio para guardar el análisis completo
            analysis_dir = Path("analisis_completo")
            analysis_dir.mkdir(exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            analysis_file = analysis_dir / f"analisis_completo_{timestamp}.json"
            
            # Guardar análisis completo
            with open(analysis_file, 'w', encoding='utf-8') as f:
                json.dump(comprehensive_data, f, indent=2, ensure_ascii=False)
            
            print(f"\n✅ Análisis completo guardado en: {analysis_file}")
            
            # Mostrar resumen del análisis
            print("\n📋 RESUMEN DEL ANÁLISIS FORENSE:")
            print("-" * 40)
            
            # Información del sistema
            system_info = comprehensive_data['analysis_info']['system_info']
            print(f"🖥️ Sistema: {system_info['system']} {system_info['release']}")
            print(f"💻 Hostname: {system_info['hostname']}")
            print(f"👤 Usuario: {system_info['username']}")
            
            # Timeline events
            timeline_count = len(comprehensive_data['timeline_events'])
            print(f"⏱️ Eventos de timeline: {timeline_count}")
            
            # Artefactos del sistema
            artifacts = comprehensive_data['system_artifacts']
            registry_count = len(artifacts.get('registry_keys', []))
            browser_count = len(artifacts.get('browser_artifacts', []))
            processes_count = len(artifacts.get('running_processes', []))
            print(f"🔑 Claves de registro: {registry_count}")
            print(f"🌐 Artefactos de navegador: {browser_count}")
            print(f"⚙️ Procesos en ejecución: {processes_count}")
            
            # Análisis de seguridad
            security = comprehensive_data['security_analysis']
            print(f"🛡️ Nivel de riesgo: {security['risk_level']}")
            print(f"📊 Puntuación de seguridad: {security['security_score']}/100")
            print(f"⚠️ Vulnerabilidades: {len(security['vulnerabilities'])}")
            print(f"🔍 Amenazas detectadas: {len(security['threats_detected'])}")
            
            # Logs de eventos
            event_logs_count = len(comprehensive_data['event_logs'])
            print(f"📋 Logs de eventos: {event_logs_count}")
            
            # Archivos Prefetch
            prefetch_count = len(comprehensive_data['prefetch_files'])
            print(f"🚀 Archivos Prefetch: {prefetch_count}")
            
            # Historial de navegadores
            browser_history_count = len(comprehensive_data['browser_history_metadata'])
            print(f"🌐 Bases de datos de navegador: {browser_history_count}")
            
            # Dispositivos USB
            usb_count = len(comprehensive_data['usb_device_history'])
            print(f"💾 Dispositivos USB detectados: {usb_count}")
            
            # Programas de inicio
            startup_count = len(comprehensive_data['startup_programs'])
            print(f"🚀 Programas de inicio: {startup_count}")
            
            # Software instalado
            software_count = len(comprehensive_data['installed_software'])
            print(f"📦 Software instalado: {software_count}")
            
            print("\n🎯 DETALLES ADICIONALES:")
            print("-" * 30)
            
            # Mostrar algunas amenazas si las hay
            if security['threats_detected']:
                print("⚠️ AMENAZAS DETECTADAS:")
                for threat in security['threats_detected'][:3]:  # Mostrar solo las primeras 3
                    print(f"   • {threat}")
            
            # Mostrar algunas vulnerabilidades
            if security['vulnerabilities']:
                print("🔍 VULNERABILIDADES:")
                for vuln in security['vulnerabilities'][:3]:  # Mostrar solo las primeras 3
                    print(f"   • {vuln}")
            
            # Mostrar algunas recomendaciones
            if security['recommendations']:
                print("💡 RECOMENDACIONES:")
                for rec in security['recommendations'][:3]:  # Mostrar solo las primeras 3
                    print(f"   • {rec}")
            
            print("\n📁 ARCHIVOS GENERADOS:")
            print(f"   📄 Análisis completo: {analysis_file}")
            
            # Preguntar si quiere ver más detalles
            print("\n🔍 ¿Deseas ver más detalles de alguna categoría?")
            print("1. 📋 Ver todos los logs de eventos")
            print("2. 🚀 Ver archivos Prefetch")
            print("3. 💾 Ver dispositivos USB")
            print("4. 📦 Ver software instalado")
            print("5. 🌐 Ver metadatos de navegadores")
            print("0. ⬅️ Volver al menú principal")
            
            detalle = input("\nSelecciona una opción: ").strip()
            
            if detalle == "1":
                self._mostrar_event_logs(comprehensive_data['event_logs'])
            elif detalle == "2":
                self._mostrar_prefetch_files(comprehensive_data['prefetch_files'])
            elif detalle == "3":
                self._mostrar_usb_devices(comprehensive_data['usb_device_history'])
            elif detalle == "4":
                self._mostrar_installed_software(comprehensive_data['installed_software'])
            elif detalle == "5":
                self._mostrar_browser_metadata(comprehensive_data['browser_history_metadata'])
            
        except Exception as e:
            print(f"❌ Error durante el análisis: {str(e)}")
            print("💡 Asegúrate de tener permisos suficientes para acceder a los recursos del sistema")
        
        input("\nPresiona Enter para continuar...")
    
    def _mostrar_event_logs(self, event_logs):
        """Mostrar detalles de logs de eventos."""
        print("\n📋 LOGS DE EVENTOS DETECTADOS:")
        print("=" * 40)
        for i, log in enumerate(event_logs[:10], 1):  # Mostrar solo los primeros 10
            if 'error' not in log:
                print(f"{i}. {log['name']}")
                print(f"   📁 Ruta: {log['path']}")
                print(f"   📊 Tamaño: {log['size_bytes']:,} bytes")
                print(f"   📅 Modificado: {log['modified']}")
                print()
        
        if len(event_logs) > 10:
            print(f"... y {len(event_logs) - 10} logs más")
    
    def _mostrar_prefetch_files(self, prefetch_files):
        """Mostrar detalles de archivos Prefetch."""
        print("\n🚀 ARCHIVOS PREFETCH DETECTADOS:")
        print("=" * 40)
        for i, pf in enumerate(prefetch_files[:10], 1):  # Mostrar solo los primeros 10
            if 'error' not in pf:
                print(f"{i}. {pf['name']}")
                print(f"   📁 Ruta: {pf['path']}")
                print(f"   📊 Tamaño: {pf['size_bytes']:,} bytes")
                print(f"   📅 Modificado: {pf['modified']}")
                print()
        
        if len(prefetch_files) > 10:
            print(f"... y {len(prefetch_files) - 10} archivos más")
    
    def _mostrar_usb_devices(self, usb_devices):
        """Mostrar detalles de dispositivos USB."""
        print("\n💾 DISPOSITIVOS USB DETECTADOS:")
        print("=" * 40)
        for i, usb in enumerate(usb_devices[:10], 1):  # Mostrar solo los primeros 10
            if 'error' not in usb:
                print(f"{i}. {usb['friendly_name']}")
                print(f"   🆔 Device ID: {usb['device_id']}")
                print(f"   🔢 Instance ID: {usb['instance_id']}")
                print()
        
        if len(usb_devices) > 10:
            print(f"... y {len(usb_devices) - 10} dispositivos más")
    
    def _mostrar_installed_software(self, installed_software):
        """Mostrar detalles de software instalado."""
        print("\n📦 SOFTWARE INSTALADO:")
        print("=" * 40)
        for i, software in enumerate(installed_software[:15], 1):  # Mostrar solo los primeros 15
            print(f"{i}. {software['name']}")
            print(f"   📊 Versión: {software['version']}")
            print(f"   🏢 Editor: {software['publisher']}")
            print(f"   📅 Instalado: {software['install_date']}")
            print()
        
        if len(installed_software) > 15:
            print(f"... y {len(installed_software) - 15} programas más")
    
    def _mostrar_browser_metadata(self, browser_metadata):
        """Mostrar metadatos de navegadores."""
        print("\n🌐 METADATOS DE NAVEGADORES:")
        print("=" * 40)
        for i, browser in enumerate(browser_metadata, 1):
            print(f"{i}. {browser['browser']}")
            print(f"   📁 Base de datos: {browser['database_path']}")
            print(f"   📊 Tamaño: {browser['size_bytes']:,} bytes")
            print(f"   📅 Modificado: {browser['modified']}")
            print()
    
    def ayuda_documentacion(self):
        """Mostrar ayuda y documentación."""
        print("\n❓ AYUDA Y DOCUMENTACIÓN")
        print("=" * 40)
        
        print("🚀 FORENSECTL - Framework de Análisis Forense Digital")
        print("\n📋 FUNCIONALIDADES PRINCIPALES:")
        print("   1. 📁 Gestión completa de casos forenses")
        print("   2. 🔍 Recopilación automática de evidencia")
        print("   3. ⚙️ Análisis forense avanzado")
        print("   4. 📄 Generación de reportes profesionales")
        print("   5. 🔗 Cadena de custodia completa")
        print("   6. 📊 Exportación de datos")
        
        print("\n🎯 FLUJO DE TRABAJO RECOMENDADO:")
        print("   1. Crear un nuevo caso")
        print("   2. Recopilar evidencia del sistema")
        print("   3. Realizar análisis forense")
        print("   4. Generar reportes")
        print("   5. Mantener cadena de custodia")
        print("   6. Exportar resultados")
        
        print("\n📚 FORMATOS SOPORTADOS:")
        print("   📄 Reportes: HTML, PDF, DOCX")
        print("   📊 Exportación: JSON, CSV")
        print("   🔍 Evidencia: JSON, TXT")
        
        print("\n🛠️ HERRAMIENTAS INCLUIDAS:")
        print("   🔍 Extractor de artefactos")
        print("   ⏱️ Constructor de timeline")
        print("   🛡️ Análisis de seguridad")
        print("   📊 Generador de estadísticas")
        
        print("\n💡 CONSEJOS:")
        print("   • Siempre mantén la cadena de custodia")
        print("   • Realiza copias de seguridad regulares")
        print("   • Documenta todos los procedimientos")
        print("   • Verifica la integridad de los datos")
        
        input("\nPresiona Enter para continuar...")
        
    def ejecutar(self):
        """Ejecutar el demo interactivo principal."""
        self.mostrar_banner()
        
        while True:
            self.mostrar_menu_principal()
            opcion = input("\nSelecciona una opción: ").strip()
            
            if opcion == "1":
                self.gestionar_casos()
            elif opcion == "2":
                self.recopilar_evidencia()
            elif opcion == "3":
                self.analisis_forense()
            elif opcion == "4":
                self.generar_reportes()
            elif opcion == "5":
                self.cadena_custodia()
            elif opcion == "6":
                self.exportacion_datos()
            elif opcion == "7":
                self.estado_sistema()
            elif opcion == "8":
                self.herramientas_adicionales()
            elif opcion == "9":
                self.analisis_forense_completo()
            elif opcion == "10":
                self.ayuda_documentacion()
            elif opcion == "0":
                print("\n👋 ¡Gracias por usar ForenseCTL!")
                print("🔍 Framework completo de análisis forense digital")
                break
            else:
                print("❌ Opción inválida. Intenta de nuevo.")


if __name__ == "__main__":
    try:
        # Instalar dependencias si es necesario
        print("🔧 Verificando dependencias...")
        
        dependencias = ['reportlab', 'python-docx']
        for dep in dependencias:
            try:
                __import__(dep.replace('-', '_'))
            except ImportError:
                print(f"📦 Instalando {dep}...")
                subprocess.run([sys.executable, "-m", "pip", "install", dep], 
                             check=True, capture_output=True)
        
        print("✅ Dependencias verificadas")
        
        # Ejecutar demo
        demo = ForenseCTLDemo()
        demo.ejecutar()
        
    except KeyboardInterrupt:
        print("\n\n⚠️ Demo interrumpido por el usuario")
    except Exception as e:
        print(f"\n❌ Error ejecutando demo: {e}")
        print("💡 Asegúrate de que ForenseCTL esté correctamente instalado")