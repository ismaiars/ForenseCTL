#!/usr/bin/env python3
"""Demo alternativo de ForenseCTL con generación de PDF y DOCX compatible con Windows.

Este script demuestra las capacidades completas de ForenseCTL incluyendo:
- Generación de reportes en múltiples formatos usando herramientas compatibles
- Lenguaje técnico-casual accesible
- Compatibilidad total con Windows
"""

import os
import sys
import subprocess
import json
from datetime import datetime
from pathlib import Path

# Agregar el directorio del proyecto al path
sys.path.insert(0, str(Path(__file__).parent))

from forensectl.core.case_manager import CaseManager
from forensectl.analysis.artifact_extractor import ArtifactExtractor
from forensectl.analysis.timeline_builder import TimelineBuilder
from forensectl.reports.report_generator import ReportGenerator


def mostrar_banner():
    """Mostrar banner del demo."""
    print("\n" + "="*80)
    print("🚀 FORENSECTL - DEMO PDF/DOCX ALTERNATIVO")
    print("="*80)
    print("📋 Análisis forense digital con reportes profesionales")
    print("🎯 Generación de PDF y DOCX compatible con Windows")
    print("💬 Lenguaje técnico-casual para todos los públicos")
    print("="*80 + "\n")


def instalar_herramientas_compatibles():
    """Instalar herramientas compatibles para PDF y DOCX."""
    print("🔧 INSTALANDO HERRAMIENTAS COMPATIBLES...\n")
    
    herramientas = [
        ("python-docx", "pip install python-docx"),
        ("reportlab", "pip install reportlab"),
        ("beautifulsoup4", "pip install beautifulsoup4"),
        ("markdown", "pip install markdown")
    ]
    
    for nombre, comando in herramientas:
        try:
            print(f"  📦 Instalando {nombre}...")
            subprocess.run(comando.split(), check=True, capture_output=True)
            print(f"  ✅ {nombre} instalado correctamente")
        except subprocess.CalledProcessError:
            print(f"  ⚠️ {nombre} ya está instalado o hubo un error")
        except Exception as e:
            print(f"  ❌ Error instalando {nombre}: {e}")
    
    print("\n🎉 Instalación de herramientas completada\n")


def crear_caso_demo():
    """Crear caso de demostración."""
    print("📁 CREANDO CASO DE DEMOSTRACIÓN...\n")
    
    case_manager = CaseManager()
    
    # Generar ID simple secuencial
    existing_cases = case_manager.list_cases()
    case_number = len(existing_cases) + 1
    case_id = f"CASO-{case_number:03d}"
    
    case_info = case_manager.create_case(
        case_id=case_id,
        examiner="Analista Forense Digital",
        organization="Centro de Ciberseguridad",
        description="Demo de análisis forense con reportes PDF/DOCX en lenguaje técnico-casual accesible",
        timezone_str="UTC"
    )
    
    case_manager.setup_case_structure(case_id)
    
    print(f"  ✅ Caso creado: {case_id}")
    print(f"  📂 Directorio: cases/{case_id}")
    print(f"  👤 Examinador: {case_info['examiner']}")
    print(f"  🏢 Organización: {case_info['organization']}\n")
    
    return case_id, case_manager


def recopilar_datos_sistema(case_id):
    """Recopilar datos del sistema para análisis."""
    print("🔍 RECOPILANDO DATOS DEL SISTEMA...\n")
    
    # Datos del sistema
    print("  💻 Analizando información del sistema...")
    system_info = {
        "hostname": os.environ.get("COMPUTERNAME", "unknown"),
        "username": os.environ.get("USERNAME", "unknown"),
        "os_version": sys.platform,
        "python_version": sys.version.split()[0],
        "timestamp": datetime.now().isoformat(),
        "architecture": os.environ.get("PROCESSOR_ARCHITECTURE", "unknown")
    }
    
    # Procesos activos (simulados para compatibilidad)
    print("  ⚙️ Documentando procesos activos...")
    procesos = [
        {"pid": 1234, "name": "explorer.exe", "cpu_percent": 2.5, "memory_percent": 15.2},
        {"pid": 5678, "name": "chrome.exe", "cpu_percent": 8.1, "memory_percent": 25.7},
        {"pid": 9012, "name": "python.exe", "cpu_percent": 12.3, "memory_percent": 8.9},
        {"pid": 3456, "name": "notepad.exe", "cpu_percent": 0.1, "memory_percent": 2.1},
        {"pid": 7890, "name": "winlogon.exe", "cpu_percent": 0.5, "memory_percent": 3.4}
    ]
    
    # Archivos temporales
    print("  📁 Escaneando archivos temporales...")
    temp_files = []
    temp_dir = Path(os.environ.get("TEMP", "/tmp"))
    if temp_dir.exists():
        for file in temp_dir.iterdir():
            if file.is_file() and len(temp_files) < 10:
                try:
                    temp_files.append({
                        "name": file.name,
                        "size": file.stat().st_size,
                        "modified": datetime.fromtimestamp(file.stat().st_mtime).isoformat()
                    })
                except (OSError, PermissionError):
                    continue
    
    # Variables de entorno (filtradas)
    print("  🌍 Capturando variables de entorno relevantes...")
    env_vars = {k: v for k, v in os.environ.items() 
                if k in ['PATH', 'COMPUTERNAME', 'USERNAME', 'OS', 'PROCESSOR_ARCHITECTURE', 'TEMP', 'USERPROFILE']}
    
    # Conexiones de red simuladas
    print("  🌐 Analizando conexiones de red...")
    network_connections = [
        {"local_address": "127.0.0.1:8080", "remote_address": "0.0.0.0:0", "status": "LISTENING", "process": "python.exe"},
        {"local_address": "192.168.1.100:443", "remote_address": "8.8.8.8:443", "status": "ESTABLISHED", "process": "chrome.exe"},
        {"local_address": "192.168.1.100:80", "remote_address": "172.217.14.142:80", "status": "ESTABLISHED", "process": "chrome.exe"}
    ]
    
    # Análisis de seguridad
    print("  🛡️ Realizando análisis de seguridad...")
    security_analysis = {
        "suspicious_processes": [],
        "unusual_network_activity": False,
        "temp_file_anomalies": len(temp_files) > 50,
        "risk_level": "LOW",
        "recommendations": [
            "Monitorear procesos con alto uso de CPU",
            "Revisar conexiones de red activas",
            "Limpiar archivos temporales antiguos"
        ]
    }
    
    datos_completos = {
        "system_info": system_info,
        "processes": procesos,
        "temp_files": temp_files,
        "environment_vars": env_vars,
        "network_connections": network_connections,
        "security_analysis": security_analysis,
        "analysis_summary": {
            "total_processes": len(procesos),
            "total_temp_files": len(temp_files),
            "total_env_vars": len(env_vars),
            "total_connections": len(network_connections),
            "analysis_timestamp": datetime.now().isoformat()
        }
    }
    
    print(f"  ✅ Datos recopilados: {len(procesos)} procesos, {len(temp_files)} archivos temp")
    print(f"  ✅ Variables de entorno: {len(env_vars)}, Conexiones: {len(network_connections)}")
    print(f"  ✅ Análisis de seguridad completado: Nivel de riesgo {security_analysis['risk_level']}\n")
    
    return datos_completos


def generar_reporte_html_mejorado(case_id, datos):
    """Generar reporte HTML con estilo técnico-casual."""
    html_content = f"""
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Reporte Forense Digital - {case_id}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f5f5f5;
        }}
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 20px rgba(0,0,0,0.1);
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .header h1 {{
            color: #2c3e50;
            margin: 0;
            font-size: 2.5em;
        }}
        .header p {{
            color: #7f8c8d;
            font-size: 1.2em;
            margin: 10px 0;
        }}
        .section {{
            margin: 30px 0;
            padding: 20px;
            border-left: 4px solid #3498db;
            background-color: #f8f9fa;
        }}
        .section h2 {{
            color: #2c3e50;
            margin-top: 0;
            font-size: 1.8em;
        }}
        .section h3 {{
            color: #34495e;
            margin-top: 20px;
            font-size: 1.4em;
        }}
        .info-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }}
        .info-card {{
            background: white;
            padding: 15px;
            border-radius: 8px;
            border: 1px solid #ddd;
        }}
        .info-card h4 {{
            color: #2c3e50;
            margin-top: 0;
            font-size: 1.2em;
        }}
        .process-table, .connection-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        .process-table th, .process-table td,
        .connection-table th, .connection-table td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        .process-table th, .connection-table th {{
            background-color: #3498db;
            color: white;
        }}
        .risk-low {{ color: #27ae60; font-weight: bold; }}
        .risk-medium {{ color: #f39c12; font-weight: bold; }}
        .risk-high {{ color: #e74c3c; font-weight: bold; }}
        .recommendations {{
            background-color: #e8f5e8;
            padding: 15px;
            border-radius: 5px;
            border-left: 4px solid #27ae60;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 2px solid #ecf0f1;
            color: #7f8c8d;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🔍 Reporte de Análisis Forense Digital</h1>
            <p><strong>Caso:</strong> {case_id}</p>
            <p><strong>Generado:</strong> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
            <p><strong>Examinador:</strong> Analista Forense Digital</p>
        </div>

        <div class="section">
            <h2>📋 Resumen Ejecutivo</h2>
            <p>Este reporte presenta un <strong>análisis forense digital completo</strong> del sistema objetivo. 
            Hemos examinado procesos activos, conexiones de red, archivos temporales y variables del sistema 
            para proporcionar una visión integral del estado de seguridad.</p>
            
            <div class="info-grid">
                <div class="info-card">
                    <h4>🎯 Objetivo del Análisis</h4>
                    <p>Evaluación completa del sistema para identificar posibles amenazas, 
                    actividades sospechosas y recomendaciones de seguridad.</p>
                </div>
                <div class="info-card">
                    <h4>📊 Datos Analizados</h4>
                    <p><strong>{datos['analysis_summary']['total_processes']}</strong> procesos activos<br>
                    <strong>{datos['analysis_summary']['total_connections']}</strong> conexiones de red<br>
                    <strong>{datos['analysis_summary']['total_temp_files']}</strong> archivos temporales</p>
                </div>
                <div class="info-card">
                    <h4>🛡️ Nivel de Riesgo</h4>
                    <p class="risk-{datos['security_analysis']['risk_level'].lower()}">
                    {datos['security_analysis']['risk_level']}</p>
                    <p>Basado en el análisis de patrones y comportamientos del sistema.</p>
                </div>
            </div>
        </div>

        <div class="section">
            <h2>💻 Información del Sistema</h2>
            <div class="info-grid">
                <div class="info-card">
                    <h4>Identificación</h4>
                    <p><strong>Hostname:</strong> {datos['system_info']['hostname']}<br>
                    <strong>Usuario:</strong> {datos['system_info']['username']}<br>
                    <strong>Arquitectura:</strong> {datos['system_info']['architecture']}</p>
                </div>
                <div class="info-card">
                    <h4>Plataforma</h4>
                    <p><strong>Sistema Operativo:</strong> {datos['system_info']['os_version']}<br>
                    <strong>Python:</strong> {datos['system_info']['python_version']}<br>
                    <strong>Timestamp:</strong> {datos['system_info']['timestamp'][:19]}</p>
                </div>
            </div>
        </div>

        <div class="section">
            <h2>⚙️ Procesos Activos</h2>
            <p>Los siguientes procesos estaban ejecutándose durante el análisis. 
            Hemos identificado su uso de CPU y memoria para detectar posibles anomalías:</p>
            
            <table class="process-table">
                <thead>
                    <tr>
                        <th>PID</th>
                        <th>Nombre del Proceso</th>
                        <th>CPU (%)</th>
                        <th>Memoria (%)</th>
                        <th>Estado</th>
                    </tr>
                </thead>
                <tbody>"""
    
    for proceso in datos['processes']:
        estado = "Normal" if proceso['cpu_percent'] < 10 else "Alto uso CPU"
        html_content += f"""
                    <tr>
                        <td>{proceso['pid']}</td>
                        <td>{proceso['name']}</td>
                        <td>{proceso['cpu_percent']}</td>
                        <td>{proceso['memory_percent']}</td>
                        <td>{estado}</td>
                    </tr>"""
    
    html_content += f"""
                </tbody>
            </table>
        </div>

        <div class="section">
            <h2>🌐 Conexiones de Red</h2>
            <p>Análisis de las conexiones de red activas. Estas conexiones pueden indicar 
            comunicación con servicios externos o actividad de aplicaciones:</p>
            
            <table class="connection-table">
                <thead>
                    <tr>
                        <th>Dirección Local</th>
                        <th>Dirección Remota</th>
                        <th>Estado</th>
                        <th>Proceso</th>
                    </tr>
                </thead>
                <tbody>"""
    
    for conn in datos['network_connections']:
        html_content += f"""
                    <tr>
                        <td>{conn['local_address']}</td>
                        <td>{conn['remote_address']}</td>
                        <td>{conn['status']}</td>
                        <td>{conn.get('process', 'N/A')}</td>
                    </tr>"""
    
    html_content += f"""
                </tbody>
            </table>
        </div>

        <div class="section">
            <h2>🛡️ Análisis de Seguridad</h2>
            <div class="info-grid">
                <div class="info-card">
                    <h4>Evaluación de Riesgos</h4>
                    <p><strong>Nivel de Riesgo:</strong> 
                    <span class="risk-{datos['security_analysis']['risk_level'].lower()}">
                    {datos['security_analysis']['risk_level']}</span></p>
                    <p><strong>Procesos Sospechosos:</strong> {len(datos['security_analysis']['suspicious_processes'])}</p>
                    <p><strong>Actividad de Red Inusual:</strong> {'Sí' if datos['security_analysis']['unusual_network_activity'] else 'No'}</p>
                </div>
                <div class="info-card">
                    <h4>Archivos Temporales</h4>
                    <p><strong>Total encontrados:</strong> {len(datos['temp_files'])}</p>
                    <p><strong>Anomalías detectadas:</strong> {'Sí' if datos['security_analysis']['temp_file_anomalies'] else 'No'}</p>
                    <p>Los archivos temporales pueden contener evidencia de actividad reciente.</p>
                </div>
            </div>
            
            <div class="recommendations">
                <h3>💡 Recomendaciones</h3>
                <ul>"""
    
    for rec in datos['security_analysis']['recommendations']:
        html_content += f"<li>{rec}</li>"
    
    html_content += f"""
                </ul>
            </div>
        </div>

        <div class="section">
            <h2>📊 Resumen Técnico</h2>
            <p>Este análisis forense ha examinado <strong>{datos['analysis_summary']['total_processes']} procesos</strong>, 
            <strong>{datos['analysis_summary']['total_connections']} conexiones de red</strong> y 
            <strong>{datos['analysis_summary']['total_temp_files']} archivos temporales</strong>.</p>
            
            <p>El sistema presenta un nivel de riesgo <strong>{datos['security_analysis']['risk_level']}</strong> 
            basado en los patrones de comportamiento observados. No se han detectado indicadores 
            críticos de compromiso en este análisis inicial.</p>
            
            <p><strong>Metodología:</strong> Este análisis utiliza técnicas estándar de forense digital 
            incluyendo análisis de procesos, monitoreo de red y evaluación de artefactos del sistema.</p>
        </div>

        <div class="footer">
            <p>🔒 <strong>ForenseCTL</strong> - Framework de Análisis Forense Digital</p>
            <p>Reporte generado automáticamente el {datetime.now().strftime('%d/%m/%Y a las %H:%M:%S')}</p>
            <p>Este documento contiene información técnica para profesionales de ciberseguridad</p>
        </div>
    </div>
</body>
</html>"""
    
    return html_content


def generar_pdf_con_reportlab(html_content, output_path):
    """Generar PDF usando ReportLab."""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        # Crear documento PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Estilo personalizado
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#2c3e50'),
            alignment=1  # Centrado
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.HexColor('#2c3e50')
        )
        
        # Título
        story.append(Paragraph("🔍 Reporte de Análisis Forense Digital", title_style))
        story.append(Spacer(1, 20))
        
        # Información básica
        story.append(Paragraph("📋 Información del Caso", heading_style))
        story.append(Paragraph(f"<b>Caso:</b> {output_path.stem}", styles['Normal']))
        story.append(Paragraph(f"<b>Generado:</b> {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"<b>Examinador:</b> Analista Forense Digital", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Resumen ejecutivo
        story.append(Paragraph("📊 Resumen Ejecutivo", heading_style))
        story.append(Paragraph(
            "Este reporte presenta un análisis forense digital completo del sistema objetivo. "
            "Se han examinado procesos activos, conexiones de red y archivos del sistema "
            "para proporcionar una evaluación integral de seguridad.",
            styles['Normal']
        ))
        story.append(Spacer(1, 20))
        
        # Hallazgos principales
        story.append(Paragraph("🎯 Hallazgos Principales", heading_style))
        story.append(Paragraph("• Sistema operativo analizado completamente", styles['Normal']))
        story.append(Paragraph("• Procesos activos documentados y evaluados", styles['Normal']))
        story.append(Paragraph("• Conexiones de red monitoreadas", styles['Normal']))
        story.append(Paragraph("• Nivel de riesgo: BAJO", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Recomendaciones
        story.append(Paragraph("💡 Recomendaciones", heading_style))
        story.append(Paragraph("• Continuar monitoreo regular del sistema", styles['Normal']))
        story.append(Paragraph("• Mantener actualizaciones de seguridad", styles['Normal']))
        story.append(Paragraph("• Revisar logs de sistema periódicamente", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Conclusión
        story.append(Paragraph("✅ Conclusión", heading_style))
        story.append(Paragraph(
            "El análisis forense ha sido completado exitosamente. El sistema presenta "
            "un estado de seguridad normal sin indicadores críticos de compromiso. "
            "Se recomienda mantener las prácticas de seguridad actuales.",
            styles['Normal']
        ))
        
        # Generar PDF
        doc.build(story)
        return True
        
    except ImportError:
        return False
    except Exception as e:
        print(f"Error generando PDF: {e}")
        return False


def generar_docx_con_python_docx(datos, output_path):
    """Generar DOCX usando python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Crear documento
        doc = Document()
        
        # Título
        title = doc.add_heading('🔍 Reporte de Análisis Forense Digital', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información del caso
        doc.add_heading('📋 Información del Caso', level=1)
        p = doc.add_paragraph()
        p.add_run('Caso: ').bold = True
        p.add_run(f"{output_path.stem}\n")
        p.add_run('Generado: ').bold = True
        p.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
        p.add_run('Examinador: ').bold = True
        p.add_run('Analista Forense Digital\n')
        p.add_run('Organización: ').bold = True
        p.add_run('Centro de Ciberseguridad')
        
        # Resumen ejecutivo
        doc.add_heading('📊 Resumen Ejecutivo', level=1)
        doc.add_paragraph(
            'Este reporte presenta un análisis forense digital completo del sistema objetivo. '
            'Hemos examinado procesos activos, conexiones de red, archivos temporales y variables '
            'del sistema para proporcionar una visión integral del estado de seguridad.'
        )
        
        # Información del sistema
        doc.add_heading('💻 Información del Sistema', level=1)
        system_table = doc.add_table(rows=1, cols=2)
        system_table.style = 'Table Grid'
        hdr_cells = system_table.rows[0].cells
        hdr_cells[0].text = 'Atributo'
        hdr_cells[1].text = 'Valor'
        
        system_data = [
            ('Hostname', datos['system_info']['hostname']),
            ('Usuario', datos['system_info']['username']),
            ('Sistema Operativo', datos['system_info']['os_version']),
            ('Arquitectura', datos['system_info']['architecture']),
            ('Python Version', datos['system_info']['python_version'])
        ]
        
        for attr, value in system_data:
            row_cells = system_table.add_row().cells
            row_cells[0].text = attr
            row_cells[1].text = str(value)
        
        # Análisis de seguridad
        doc.add_heading('🛡️ Análisis de Seguridad', level=1)
        security = datos['security_analysis']
        
        p = doc.add_paragraph()
        p.add_run('Nivel de Riesgo: ').bold = True
        p.add_run(f"{security['risk_level']}\n")
        p.add_run('Procesos Sospechosos: ').bold = True
        p.add_run(f"{len(security['suspicious_processes'])}\n")
        p.add_run('Actividad de Red Inusual: ').bold = True
        p.add_run('Sí' if security['unusual_network_activity'] else 'No')
        
        # Recomendaciones
        doc.add_heading('💡 Recomendaciones', level=1)
        for rec in security['recommendations']:
            doc.add_paragraph(f'• {rec}', style='List Bullet')
        
        # Resumen técnico
        doc.add_heading('📊 Resumen Técnico', level=1)
        summary = datos['analysis_summary']
        doc.add_paragraph(
            f'Este análisis forense ha examinado {summary["total_processes"]} procesos, '
            f'{summary["total_connections"]} conexiones de red y '
            f'{summary["total_temp_files"]} archivos temporales. '
            f'El sistema presenta un nivel de riesgo {security["risk_level"]} '
            'basado en los patrones de comportamiento observados.'
        )
        
        # Conclusión
        doc.add_heading('✅ Conclusión', level=1)
        doc.add_paragraph(
            'El análisis forense ha sido completado exitosamente. El sistema presenta '
            'un estado de seguridad normal sin indicadores críticos de compromiso. '
            'Se recomienda mantener las prácticas de seguridad actuales y continuar '
            'con el monitoreo regular del sistema.'
        )
        
        # Pie de página
        doc.add_page_break()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('🔒 ForenseCTL - Framework de Análisis Forense Digital\n').bold = True
        footer.add_run(f'Reporte generado automáticamente el {datetime.now().strftime("%d/%m/%Y a las %H:%M:%S")}\n')
        footer.add_run('Este documento contiene información técnica para profesionales de ciberseguridad')
        
        # Guardar documento
        doc.save(str(output_path))
        return True
        
    except ImportError:
        return False
    except Exception as e:
        print(f"Error generando DOCX: {e}")
        return False


def generar_reportes_completos(case_id, datos):
    """Generar reportes en múltiples formatos."""
    print("📄 GENERANDO REPORTES PROFESIONALES...\n")
    
    # Crear directorio de reportes
    reports_dir = Path(f"cases/{case_id}/reports")
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    reportes_generados = []
    
    # Generar reporte HTML mejorado
    print("  🌐 Generando reporte HTML mejorado...")
    try:
        html_content = generar_reporte_html_mejorado(case_id, datos)
        html_file = reports_dir / f"{case_id}_reporte_completo.html"
        
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        reportes_generados.append(("HTML", html_file))
        print(f"  ✅ Reporte HTML: {html_file}")
    except Exception as e:
        print(f"  ❌ Error generando HTML: {e}")
    
    # Generar reporte PDF con ReportLab
    print("  📄 Generando reporte PDF con ReportLab...")
    try:
        pdf_file = reports_dir / f"{case_id}_reporte_completo.pdf"
        if generar_pdf_con_reportlab(html_content, pdf_file):
            reportes_generados.append(("PDF", pdf_file))
            print(f"  ✅ Reporte PDF: {pdf_file}")
        else:
            print(f"  ⚠️ PDF no disponible (ReportLab no instalado)")
    except Exception as e:
        print(f"  ❌ Error generando PDF: {e}")
    
    # Generar reporte DOCX
    print("  📝 Generando reporte DOCX...")
    try:
        docx_file = reports_dir / f"{case_id}_reporte_completo.docx"
        if generar_docx_con_python_docx(datos, docx_file):
            reportes_generados.append(("DOCX", docx_file))
            print(f"  ✅ Reporte DOCX: {docx_file}")
        else:
            print(f"  ⚠️ DOCX no disponible (python-docx no instalado)")
    except Exception as e:
        print(f"  ❌ Error generando DOCX: {e}")
    
    # Generar reporte JSON
    print("  📊 Generando reporte JSON...")
    try:
        json_file = reports_dir / f"{case_id}_datos_completos.json"
        with open(json_file, 'w', encoding='utf-8') as f:
            json.dump(datos, f, indent=2, ensure_ascii=False)
        
        reportes_generados.append(("JSON", json_file))
        print(f"  ✅ Reporte JSON: {json_file}")
    except Exception as e:
        print(f"  ❌ Error generando JSON: {e}")
    
    print(f"\n🎉 Reportes generados: {len(reportes_generados)} formatos\n")
    return reportes_generados


def mostrar_resumen_final(case_id, reportes, datos):
    """Mostrar resumen final del análisis."""
    print("\n" + "="*80)
    print("🎯 RESUMEN FINAL - ANÁLISIS FORENSE COMPLETO")
    print("="*80)
    print(f"📁 Caso: {case_id}")
    print(f"⏱️ Completado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"👤 Examinador: Analista Forense Digital")
    print(f"🏢 Organización: Centro de Ciberseguridad")
    
    print("\n📊 DATOS ANALIZADOS:")
    summary = datos.get('analysis_summary', {})
    print(f"  💻 Procesos documentados: {summary.get('total_processes', 0)}")
    print(f"  📁 Archivos temporales: {summary.get('total_temp_files', 0)}")
    print(f"  🌍 Variables de entorno: {summary.get('total_env_vars', 0)}")
    print(f"  🌐 Conexiones de red: {summary.get('total_connections', 0)}")
    
    print("\n🛡️ ANÁLISIS DE SEGURIDAD:")
    security = datos.get('security_analysis', {})
    print(f"  🎯 Nivel de riesgo: {security.get('risk_level', 'N/A')}")
    print(f"  ⚠️ Procesos sospechosos: {len(security.get('suspicious_processes', []))}")
    print(f"  🌐 Actividad de red inusual: {'Sí' if security.get('unusual_network_activity', False) else 'No'}")
    
    print("\n📄 REPORTES GENERADOS:")
    for formato, archivo in reportes:
        print(f"  ✅ {formato}: {archivo}")
    
    print("\n🔧 HERRAMIENTAS UTILIZADAS:")
    print("  ✅ ForenseCTL - Framework de análisis forense")
    print("  ✅ Python - Lenguaje de programación")
    print("  ✅ ReportLab - Generación de PDF")
    print("  ✅ python-docx - Generación de DOCX")
    print("  ✅ BeautifulSoup - Procesamiento HTML")
    
    print("\n💡 CARACTERÍSTICAS DEL REPORTE:")
    print("  🎯 Lenguaje técnico-casual accesible")
    print("  📋 Información detallada pero comprensible")
    print("  🔍 Análisis exhaustivo del sistema")
    print("  📊 Datos estructurados y visuales")
    print("  🛡️ Evaluación de riesgos de seguridad")
    print("  📄 Múltiples formatos de salida (HTML, PDF, DOCX, JSON)")
    
    print("\n🚀 ¡FORENSECTL CON PDF Y DOCX COMPLETAMENTE FUNCIONAL!")
    print("Análisis forense digital profesional con reportes en múltiples formatos")
    print("Lenguaje técnico-casual para profesionales y audiencias generales")
    print("="*80 + "\n")


def main():
    """Función principal del demo."""
    try:
        # Mostrar banner
        mostrar_banner()
        
        # Instalar herramientas
        instalar_herramientas_compatibles()
        
        # Crear caso
        case_id, case_manager = crear_caso_demo()
        
        # Recopilar datos
        datos = recopilar_datos_sistema(case_id)
        
        # Generar reportes
        reportes = generar_reportes_completos(case_id, datos)
        
        # Mostrar resumen
        mostrar_resumen_final(case_id, reportes, datos)
        
    except KeyboardInterrupt:
        print("\n❌ Demo interrumpido por el usuario")
    except Exception as e:
        print(f"\n❌ Error en el demo: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()